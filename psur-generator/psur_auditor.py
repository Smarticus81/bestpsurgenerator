"""LLM-powered PSUR Auditor — audits .docx PSURs and in-memory PSUR JSON dicts
against EU MDR MDCG 2022-21 and UK MDR requirements.

Both EU MDR (MDCG 2022-21) and UK MDR requirements are ALWAYS applied. This
mirrors regulatory reality: any device marketed in both jurisdictions must
satisfy both frameworks simultaneously.

Two-pass architecture:
  1. **Keyword pre-screen** — fast regex / keyword checks for each requirement.
  2. **LLM deep-analysis** — sends section text + requirement description to the
     LLM for nuanced compliance assessment. Batched per PSUR section to minimise
     API calls.

Integration modes:
  - **Pipeline audit** (``run_json_audit``): Audits in-memory PSUR JSON during
    generation and returns per-section remediation instructions so the
    orchestrator can fix gaps iteratively.
  - **Post-hoc DOCX audit** (``run_audit``): Audits a rendered .docx file.
  - **Standalone CLI**: ``python psur_auditor.py path/to/PSUR.docx``
"""

from __future__ import annotations

import json
import re
import sys
import time
from dataclasses import asdict, dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document  # python-docx
from rich.console import Console
from rich.panel import Panel
from rich.table import Table

# ---------------------------------------------------------------------------
# Project imports
# ---------------------------------------------------------------------------
from config import MODEL
from llm_client import create_message

console = Console()

# ═══════════════════════════════════════════════════════════════════════════
# Data models
# ═══════════════════════════════════════════════════════════════════════════


class ComplianceStatus(str, Enum):
    """Traffic-light compliance status."""
    COMPLIANT = "COMPLIANT"
    PARTIAL = "PARTIAL"
    GAP = "GAP"
    NOT_APPLICABLE = "NOT_APPLICABLE"


class Framework(str, Enum):
    EU_MDR = "EU_MDR"
    UK_MDR = "UK_MDR"
    BOTH = "BOTH"


class Criticality(str, Enum):
    CRITICAL = "CRITICAL"
    MAJOR = "MAJOR"
    MINOR = "MINOR"


@dataclass
class Requirement:
    """Single auditable requirement."""
    req_id: str
    title: str
    description: str
    section: str              # PSUR section letter (A–M) or meta
    framework: Framework
    criticality: Criticality
    keywords: List[str]       # fast pre-screen keywords
    guidance_ref: str = ""    # e.g. "MDCG 2022-21 Annex I"


@dataclass
class AuditFinding:
    """Result for one requirement."""
    req_id: str
    title: str
    status: ComplianceStatus
    evidence: str             # quote or summary supporting status
    recommendation: str       # what to fix
    section: str
    framework: str
    criticality: str
    llm_assessed: bool = False


@dataclass
class AuditReport:
    """Full audit output."""
    psur_path: str
    audit_timestamp: str
    total_requirements: int = 0
    compliant: int = 0
    partial: int = 0
    gap: int = 0
    not_applicable: int = 0
    compliance_score: float = 0.0
    findings: List[AuditFinding] = field(default_factory=list)
    llm_summary: str = ""
    uk_mdr_enabled: bool = False
    llm_enabled: bool = True
    token_usage: Dict[str, int] = field(default_factory=lambda: {"input": 0, "output": 0})

    def compute_score(self) -> None:
        self.compliant = sum(1 for f in self.findings if f.status == ComplianceStatus.COMPLIANT)
        self.partial = sum(1 for f in self.findings if f.status == ComplianceStatus.PARTIAL)
        self.gap = sum(1 for f in self.findings if f.status == ComplianceStatus.GAP)
        self.not_applicable = sum(1 for f in self.findings if f.status == ComplianceStatus.NOT_APPLICABLE)
        self.total_requirements = len(self.findings)
        scorable = self.total_requirements - self.not_applicable
        if scorable > 0:
            self.compliance_score = round(
                (self.compliant + 0.5 * self.partial) / scorable * 100, 1
            )


# ═══════════════════════════════════════════════════════════════════════════
# Requirements checklist builder
# ═══════════════════════════════════════════════════════════════════════════

def build_requirements_checklist(*, include_uk: bool = True) -> List[Requirement]:
    """Build the full requirements checklist from constraints/audit_requirements.json.

    Includes EU MDR (always) and UK MDR (when include_uk=True). The JSON file is the
    authoritative source — edit it to add, remove, or modify auditable requirements
    without touching code.
    """
    json_path = Path(__file__).parent / "constraints" / "audit_requirements.json"
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    reqs: List[Requirement] = []
    for r in data["requirements"]:
        framework = Framework(r["framework"])
        if framework == Framework.UK_MDR and not include_uk:
            continue
        reqs.append(Requirement(
            req_id=r["req_id"],
            title=r["title"],
            description=r["description"],
            section=r["section"],
            framework=framework,
            criticality=Criticality(r["criticality"]),
            keywords=list(r.get("keywords") or []),
            guidance_ref=r.get("guidance_ref", ""),
        ))
    return reqs


# ═══════════════════════════════════════════════════════════════════════════
# PSUR Document Parser
# ═══════════════════════════════════════════════════════════════════════════

# Map PSUR section letters to likely heading patterns
_SECTION_HEADING_PATTERNS: Dict[str, List[str]] = {
    "A": [r"executive\s+summary"],
    "B": [r"scope", r"device\s+description", r"description\s+of.*device",
          r"intended\s+(purpose|use)"],
    "C": [r"volume\s+of\s+(sales|distribution)", r"population\s+exposure",
          r"sales\s+and\s+population"],
    "D": [r"serious\s+incident", r"vigilance", r"information\s+on\s+serious"],
    "E": [r"customer\s+feedback", r"user\s+feedback", r"feedback\s+and\s+complaint"],
    "F": [r"complaint\s+type", r"product\s+complaint", r"complaint\s+count",
          r"complaint\s+rate"],
    "G": [r"trend\s+report", r"trend\s+analysis", r"information\s+from\s+trend"],
    "H": [r"field\s+safety", r"fsca"],
    "I": [r"corrective\s+and\s+preventive", r"capa"],
    "J": [r"scientific\s+literature", r"literature\s+review"],
    "K": [r"external\s+database", r"registr", r"publicly\s+available"],
    "L": [r"post-?market\s+clinical", r"pmcf"],
    "M": [r"finding.*conclusion", r"overall\s+conclusion", r"benefit.*risk\s+(?:determination|conclusion)"],
}


class PSURParser:
    """Extract structured content from a PSUR .docx file."""

    def __init__(self, docx_path: str | Path) -> None:
        self.path = Path(docx_path)
        self.doc = Document(str(self.path))
        self.full_text = ""
        self.paragraphs: List[str] = []
        self.headings: List[Tuple[str, int]] = []  # (text, level)
        self.sections: Dict[str, str] = {}          # section_key → text
        self._parse()

    def _parse(self) -> None:
        """Extract paragraphs, headings, and build section map."""
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            self.paragraphs.append(text)
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                try:
                    level = int(re.search(r"\d", style_name).group())
                except (AttributeError, ValueError):
                    level = 1
                self.headings.append((text, level))

        self.full_text = "\n".join(self.paragraphs)

        # Build section map keyed by PSUR section letter
        self._map_sections()

    def _map_sections(self) -> None:
        """Map content to PSUR sections A–M based on heading patterns."""
        # Combine all paragraph text with indices
        para_texts = self.paragraphs
        total = len(para_texts)

        # Find heading indices that match section patterns
        heading_positions: List[Tuple[str, int]] = []  # (section_key, para_index)
        for i, text in enumerate(para_texts):
            text_lower = text.lower()
            for sec_key, patterns in _SECTION_HEADING_PATTERNS.items():
                if any(re.search(p, text_lower) for p in patterns):
                    heading_positions.append((sec_key, i))
                    break

        # Sort by position
        heading_positions.sort(key=lambda x: x[1])

        # Extract text between heading positions
        for idx, (sec_key, start) in enumerate(heading_positions):
            end = heading_positions[idx + 1][1] if idx + 1 < len(heading_positions) else total
            section_text = "\n".join(para_texts[start:end])
            # If multiple headings map to same section, concatenate
            if sec_key in self.sections:
                self.sections[sec_key] += "\n" + section_text
            else:
                self.sections[sec_key] = section_text

    def get_section(self, key: str) -> str:
        """Get text for a PSUR section by letter (A–M)."""
        return self.sections.get(key.upper(), "")

    def search_text(self, pattern: str, case_sensitive: bool = False) -> List[str]:
        """Find all paragraphs matching a regex pattern."""
        flags = 0 if case_sensitive else re.IGNORECASE
        return [p for p in self.paragraphs if re.search(pattern, p, flags)]

    def has_tables(self) -> bool:
        return len(self.doc.tables) > 0

    def get_table_count(self) -> int:
        return len(self.doc.tables)

    def get_word_count(self) -> int:
        return len(self.full_text.split())

    def get_section_count(self) -> int:
        return len(self.sections)


# ═══════════════════════════════════════════════════════════════════════════
# Keyword pre-screen engine
# ═══════════════════════════════════════════════════════════════════════════

def _keyword_prescreen(
    req: Requirement,
    parser: PSURParser,
) -> Tuple[ComplianceStatus, str]:
    """Fast keyword/regex check. Returns (status, evidence_snippet)."""

    # --- Special structural checks ---
    if req.req_id == "R.01":
        # Check for prohibited regulation citations
        bad_patterns = [r"MDR\s+Article\s+\d+", r"MDCG\s+\d{4}"]
        hits = []
        for bp in bad_patterns:
            hits.extend(parser.search_text(bp))
        if hits:
            return ComplianceStatus.GAP, f"Prohibited citations found: {hits[0][:120]}"
        return ComplianceStatus.COMPLIANT, "No prohibited regulation citations detected."

    if req.req_id == "R.02":
        # Check for bullet points (lines starting with •, -, *)
        bullet_lines = [p for p in parser.paragraphs if re.match(r"^\s*[•\-\*]\s", p)]
        if len(bullet_lines) > 5:
            return ComplianceStatus.GAP, f"Found {len(bullet_lines)} bullet-point lines."
        elif bullet_lines:
            return ComplianceStatus.PARTIAL, f"Found {len(bullet_lines)} bullet-point lines."
        return ComplianceStatus.COMPLIANT, "No prohibited bullet points detected."

    if not req.keywords:
        return ComplianceStatus.PARTIAL, "No keywords defined for pre-screen; needs LLM analysis."

    # --- Normal keyword check ---
    # Determine search scope: section-specific or full document
    section_text = ""
    if req.section in _SECTION_HEADING_PATTERNS:
        section_text = parser.get_section(req.section)

    search_text = section_text if section_text else parser.full_text

    if not search_text:
        return ComplianceStatus.GAP, "Section not found in document."

    search_lower = search_text.lower()
    matched_kw = [kw for kw in req.keywords if kw.lower() in search_lower]
    match_ratio = len(matched_kw) / len(req.keywords) if req.keywords else 0

    if match_ratio >= 0.6:
        # Find first matching snippet
        for kw in matched_kw:
            idx = search_lower.find(kw.lower())
            if idx >= 0:
                start = max(0, idx - 60)
                end = min(len(search_text), idx + len(kw) + 60)
                snippet = search_text[start:end].replace("\n", " ")
                return ComplianceStatus.COMPLIANT, f"...{snippet}..."
        return ComplianceStatus.COMPLIANT, f"Keywords found: {', '.join(matched_kw)}"
    elif match_ratio > 0:
        snippet = f"Partial match ({len(matched_kw)}/{len(req.keywords)} keywords): {', '.join(matched_kw)}"
        return ComplianceStatus.PARTIAL, snippet
    else:
        return ComplianceStatus.GAP, f"None of the expected keywords found: {', '.join(req.keywords[:5])}"


# ═══════════════════════════════════════════════════════════════════════════
# LLM deep-analysis engine
# ═══════════════════════════════════════════════════════════════════════════

_LLM_SYSTEM_PROMPT = """\
You are an expert medical device regulatory auditor specialising in EU MDR \
(Regulation (EU) 2017/745) and MDCG 2022-21 guidance on PSURs. You also have \
deep knowledge of UK MDR post-market surveillance requirements.

You are auditing a PSUR (Periodic Safety Update Report) document. For each \
requirement presented, you will receive:
- The requirement ID, title, description, and regulatory reference
- The relevant section text extracted from the PSUR document

Your task: Assess whether the PSUR text satisfies the requirement. Return a \
JSON object with EXACTLY these fields:
{
  "status": "COMPLIANT" | "PARTIAL" | "GAP" | "NOT_APPLICABLE",
  "evidence": "<specific quote or summary from the text supporting your assessment — max 200 chars>",
  "recommendation": "<what the manufacturer should fix/add — max 200 chars, empty string if COMPLIANT>"
}

Rules:
- COMPLIANT: The requirement is fully addressed with adequate detail.
- PARTIAL: Some relevant content exists but is incomplete or lacks depth.
- GAP: The requirement is not addressed or critically missing.
- NOT_APPLICABLE: The requirement does not apply to the document context (e.g. no UK sales for UK requirements).
- Be precise. Quote specific text as evidence when possible.
- If the section text is empty, return GAP with evidence "Section not found in document."
- Return ONLY the JSON object, no other text.
"""


def _build_llm_audit_batch(
    requirements: List[Tuple[Requirement, ComplianceStatus, str]],
    parser: PSURParser,
    section_key: str,
) -> Optional[str]:
    """Build a single LLM prompt for auditing a batch of requirements for one section.

    Returns None if no requirements need LLM analysis.
    """
    # Only send requirements that are PARTIAL or GAP from keyword pre-screen
    needs_llm = [(r, kw_status, kw_evidence)
                 for r, kw_status, kw_evidence in requirements
                 if kw_status in (ComplianceStatus.PARTIAL, ComplianceStatus.GAP)]

    if not needs_llm:
        return None

    # Get section text (truncated for token budget)
    section_text = parser.get_section(section_key)
    if not section_text and section_key not in _SECTION_HEADING_PATTERNS:
        # For meta-sections (N, O, P, Q, R, S), use full text (truncated)
        section_text = parser.full_text

    if not section_text:
        section_text = "(Section not found in the document)"

    # Truncate to ~8000 chars to stay within token budget
    max_chars = 8000
    if len(section_text) > max_chars:
        section_text = section_text[:max_chars] + "\n... [truncated]"

    # Build requirement list
    req_items = []
    for r, kw_status, kw_evidence in needs_llm:
        req_items.append(
            f"- **{r.req_id}** ({r.title}): {r.description}\n"
            f"  Reference: {r.guidance_ref}\n"
            f"  Pre-screen result: {kw_status.value} — {kw_evidence}"
        )

    prompt = (
        f"## PSUR Section Text\n\n{section_text}\n\n"
        f"---\n\n"
        f"## Requirements to Audit ({len(needs_llm)})\n\n"
        + "\n".join(req_items)
        + "\n\n---\n\n"
        f"For EACH requirement above, provide your assessment as a JSON array "
        f"of objects (one per requirement, in order). Each object must have "
        f"exactly: {{\"req_id\": \"...\", \"status\": \"...\", \"evidence\": \"...\", \"recommendation\": \"...\"}}\n\n"
        f"Return ONLY the JSON array."
    )
    return prompt


def _parse_llm_response(
    response_text: str,
    requirement_ids: List[str],
) -> Dict[str, Dict[str, str]]:
    """Parse LLM response JSON array into a dict keyed by req_id."""
    results: Dict[str, Dict[str, str]] = {}

    # Extract JSON array using bracket matching
    text = response_text.strip()
    start = text.find("[")
    if start == -1:
        # Try to parse as single object
        start = text.find("{")
        if start == -1:
            return results
        # Wrap in array
        text = "[" + text[start:] + "]"
    else:
        text = text[start:]

    # Find matching closing bracket
    depth = 0
    end = -1
    for i, ch in enumerate(text):
        if ch == "[":
            depth += 1
        elif ch == "]":
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    if end > 0:
        text = text[:end]

    try:
        items = json.loads(text)
        if isinstance(items, list):
            for item in items:
                rid = item.get("req_id", "")
                if rid:
                    results[rid] = {
                        "status": item.get("status", "PARTIAL"),
                        "evidence": item.get("evidence", ""),
                        "recommendation": item.get("recommendation", ""),
                    }
        elif isinstance(items, dict):
            rid = items.get("req_id", requirement_ids[0] if requirement_ids else "")
            if rid:
                results[rid] = {
                    "status": items.get("status", "PARTIAL"),
                    "evidence": items.get("evidence", ""),
                    "recommendation": items.get("recommendation", ""),
                }
    except json.JSONDecodeError:
        pass

    return results


def _generate_audit_summary(
    report: AuditReport,
    parser: PSURParser,
) -> str:
    """Generate an LLM-powered executive summary of the audit findings."""
    # Build a concise summary of findings for the LLM
    gap_findings = [f for f in report.findings if f.status == ComplianceStatus.GAP]
    partial_findings = [f for f in report.findings if f.status == ComplianceStatus.PARTIAL]

    findings_text = ""
    if gap_findings:
        findings_text += "### Critical Gaps:\n"
        for f in gap_findings:
            findings_text += f"- {f.req_id} ({f.title}): {f.evidence}\n"
    if partial_findings:
        findings_text += "\n### Partial Compliance:\n"
        for f in partial_findings[:10]:  # Limit for token budget
            findings_text += f"- {f.req_id} ({f.title}): {f.evidence}\n"

    prompt = (
        f"You are a senior regulatory auditor writing the executive summary for a "
        f"PSUR compliance audit report.\n\n"
        f"## Audit Statistics\n"
        f"- Total requirements: {report.total_requirements}\n"
        f"- Compliant: {report.compliant}\n"
        f"- Partial: {report.partial}\n"
        f"- Gap: {report.gap}\n"
        f"- Not applicable: {report.not_applicable}\n"
        f"- Compliance score: {report.compliance_score}%\n"
        f"- UK MDR scope: {'Yes' if report.uk_mdr_enabled else 'No'}\n"
        f"- Document word count: {parser.get_word_count()}\n"
        f"- Sections identified: {parser.get_section_count()}\n"
        f"- Tables present: {parser.get_table_count()}\n\n"
        f"{findings_text}\n\n"
        f"Write a concise 3-5 paragraph executive summary of the audit findings. "
        f"Highlight the most critical gaps, areas of strength, and priority "
        f"recommendations. Use a professional regulatory tone. Do NOT use bullet "
        f"points. Return ONLY the narrative text."
    )

    try:
        resp = create_message(
            model=MODEL,
            max_tokens=1500,
            temperature=0.2,
            system="You are a senior medical device regulatory auditor.",
            messages=[{"role": "user", "content": prompt}],
        )
        report.token_usage["input"] += resp.usage.input_tokens
        report.token_usage["output"] += resp.usage.output_tokens
        return resp.content[0].text.strip()
    except Exception as e:
        return f"(LLM summary generation failed: {e})"


# ═══════════════════════════════════════════════════════════════════════════
# In-pipeline JSON audit engine  (for the generation loop)
# ═══════════════════════════════════════════════════════════════════════════

# Maps PSUR section keys (A_executive_summary, etc.) → requirement section
# letters used in the requirements checklist, PLUS cross-section meta-reqs.
_SECTION_KEY_TO_LETTERS: Dict[str, List[str]] = {
    "A_executive_summary":                           ["A"],
    "B_scope_and_device_description":                ["B"],
    "C_volume_of_sales_and_population_exposure":     ["C"],
    "D_information_on_serious_incidents":             ["D"],
    "E_customer_feedback":                           ["E"],
    "F_product_complaint_types_counts_and_rates":    ["F"],
    "G_information_from_trend_reporting":             ["G"],
    "H_information_from_fsca":                       ["H"],
    "I_corrective_and_preventive_actions":            ["I"],
    "J_scientific_literature_review":                 ["J"],
    "K_review_of_external_databases_and_registries":  ["K"],
    "L_pmcf":                                        ["L"],
    "M_findings_and_conclusions":                     ["M"],
}

# Requirements whose section letters map to meta-categories (N-S) are checked
# against the full PSUR only once, not per-section.
_META_SECTION_LETTERS = {"N", "O", "P", "Q", "R", "S"}


def _flatten_json_to_text(obj: Any, depth: int = 0) -> str:
    """Recursively flatten a JSON section dict into searchable text."""
    parts: List[str] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            if k.startswith("_"):
                continue
            child_text = _flatten_json_to_text(v, depth + 1)
            if child_text.strip():
                parts.append(child_text)
    elif isinstance(obj, list):
        for item in obj:
            parts.append(_flatten_json_to_text(item, depth + 1))
    elif isinstance(obj, str) and obj.strip():
        parts.append(obj)
    elif isinstance(obj, (int, float)):
        parts.append(str(obj))
    return "\n".join(parts)


def _keyword_prescreen_json(
    req: Requirement,
    section_text: str,
    full_text: str,
) -> Tuple[ComplianceStatus, str]:
    """Keyword pre-screen against flattened JSON text (no PSURParser needed)."""
    # Special structural checks
    if req.req_id == "R.01":
        bad_patterns = [r"MDR\s+Article\s+\d+", r"MDCG\s+\d{4}"]
        hits = [m.group() for bp in bad_patterns
                for m in re.finditer(bp, full_text, re.IGNORECASE)]
        if hits:
            return ComplianceStatus.GAP, f"Prohibited citations found: {hits[0][:120]}"
        return ComplianceStatus.COMPLIANT, "No prohibited regulation citations detected."

    if req.req_id == "R.02":
        bullet_lines = re.findall(r"^\s*[•\-\*]\s", full_text, re.MULTILINE)
        if len(bullet_lines) > 5:
            return ComplianceStatus.GAP, f"Found {len(bullet_lines)} bullet-point lines."
        elif bullet_lines:
            return ComplianceStatus.PARTIAL, f"Found {len(bullet_lines)} bullet-point lines."
        return ComplianceStatus.COMPLIANT, "No prohibited bullet points detected."

    if not req.keywords:
        return ComplianceStatus.PARTIAL, "No keywords for pre-screen; needs LLM analysis."

    # Use section text when available, fall back to full text
    search_text = section_text if section_text else full_text
    if not search_text:
        return ComplianceStatus.GAP, "Section not found in PSUR."

    search_lower = search_text.lower()
    matched_kw = [kw for kw in req.keywords if kw.lower() in search_lower]
    match_ratio = len(matched_kw) / len(req.keywords) if req.keywords else 0

    if match_ratio >= 0.6:
        for kw in matched_kw:
            idx = search_lower.find(kw.lower())
            if idx >= 0:
                start = max(0, idx - 60)
                end = min(len(search_text), idx + len(kw) + 60)
                snippet = search_text[start:end].replace("\n", " ")
                return ComplianceStatus.COMPLIANT, f"...{snippet}..."
        return ComplianceStatus.COMPLIANT, f"Keywords found: {', '.join(matched_kw)}"
    elif match_ratio > 0:
        return ComplianceStatus.PARTIAL, (
            f"Partial ({len(matched_kw)}/{len(req.keywords)} keywords): {', '.join(matched_kw)}"
        )
    else:
        return ComplianceStatus.GAP, (
            f"None of the expected keywords found: {', '.join(req.keywords[:5])}"
        )


def _build_llm_json_audit_batch(
    requirements: List[Tuple[Requirement, ComplianceStatus, str]],
    section_text: str,
) -> Optional[str]:
    """Build LLM prompt for auditing a batch of reqs against JSON-derived text."""
    needs_llm = [(r, s, e) for r, s, e in requirements
                 if s in (ComplianceStatus.PARTIAL, ComplianceStatus.GAP)]
    if not needs_llm:
        return None

    if not section_text:
        section_text = "(Section not found in the PSUR)"

    max_chars = 10000
    if len(section_text) > max_chars:
        section_text = section_text[:max_chars] + "\n... [truncated]"

    req_items = []
    for r, kw_status, kw_evidence in needs_llm:
        req_items.append(
            f"- **{r.req_id}** ({r.title}): {r.description}\n"
            f"  Reference: {r.guidance_ref}\n"
            f"  Pre-screen: {kw_status.value} — {kw_evidence}"
        )

    return (
        f"## PSUR Section Content\n\n{section_text}\n\n---\n\n"
        f"## Requirements to Audit ({len(needs_llm)})\n\n"
        + "\n".join(req_items)
        + "\n\n---\n\n"
        f"For EACH requirement above, return a JSON array of objects "
        f"(one per requirement, in order). Each object must have exactly: "
        f'{{\"req_id\": \"...\", \"status\": \"...\", \"evidence\": \"...\", \"recommendation\": \"...\"}}\n\n'
        f"Return ONLY the JSON array."
    )


@dataclass
class SectionAuditResult:
    """Audit result for a single PSUR section, used by the remediation loop."""
    section_key: str
    findings: List[AuditFinding]
    has_gaps: bool = False
    has_critical_gaps: bool = False
    remediation_prompt: str = ""

    def build_remediation_prompt(self) -> str:
        """Build a focused remediation prompt from gap/partial findings."""
        actionable = [f for f in self.findings
                      if f.status in (ComplianceStatus.GAP, ComplianceStatus.PARTIAL)]
        if not actionable:
            self.remediation_prompt = ""
            return ""

        lines = [
            "## COMPLIANCE AUDIT FINDINGS — MANDATORY REMEDIATION\n",
            "The following compliance gaps were identified by the PSUR auditor. "
            "You MUST address ALL of them in your revised output.\n",
        ]

        for f in actionable:
            severity = "CRITICAL GAP" if f.criticality == "CRITICAL" else (
                "MAJOR GAP" if f.status == ComplianceStatus.GAP else "PARTIAL"
            )
            lines.append(
                f"### [{severity}] {f.req_id}: {f.title}\n"
                f"- Framework: {f.framework}\n"
                f"- Finding: {f.evidence}\n"
                f"- Required action: {f.recommendation}\n"
            )

        lines.append(
            "\n## INSTRUCTIONS\n"
            "1. Read each finding above carefully.\n"
            "2. Ensure your revised JSON output addresses EVERY finding.\n"
            "3. For GAP findings: add the missing content entirely.\n"
            "4. For PARTIAL findings: expand and deepen the existing content.\n"
            "5. Do NOT remove any existing valid content.\n"
            "6. Output the COMPLETE section JSON.\n"
        )

        self.remediation_prompt = "\n".join(lines)
        return self.remediation_prompt


def run_json_audit(
    psur: Dict[str, Any],
    *,
    uk_market_detected: bool = False,
    use_llm: bool = True,
    verbose: bool = False,
) -> Tuple[List[SectionAuditResult], AuditReport]:
    """Audit an in-memory PSUR JSON dict during generation.

    This is the pipeline-integrated audit. It always applies both EU MDR
    (MDCG 2022-21) and UK MDR requirements:
      - EU MDR requirements always apply.
      - UK MDR requirements apply when uk_market_detected is True; otherwise
        they are assessed but marked NOT_APPLICABLE.

    Args:
        psur: The PSUR dict (with psur["sections"] containing A–M).
        uk_market_detected: Whether UK sales exist.
        use_llm: Whether to invoke LLM for deep analysis (True by default).
        verbose: Console output.

    Returns:
        Tuple of:
          - List of SectionAuditResult (one per section, with remediation prompts)
          - Full AuditReport for logging/reporting
    """
    t0 = time.time()

    # Always include UK requirements — they'll be NOT_APPLICABLE if no UK sales
    requirements = build_requirements_checklist(include_uk=True)
    sections = psur.get("sections", {})

    # Pre-flatten all section texts
    section_texts: Dict[str, str] = {}
    for sec_key, sec_data in sections.items():
        section_texts[sec_key] = _flatten_json_to_text(sec_data)

    full_text = "\n".join(section_texts.values())

    if verbose:
        console.print(f"\n[bold]  Auditing PSUR ({len(requirements)} requirements, "
                       f"{len(sections)} sections)...[/bold]")

    # ── 1. Keyword pre-screen ─────────────────────────────────────────
    keyword_results: Dict[str, Tuple[ComplianceStatus, str]] = {}
    for req in requirements:
        # UK requirements → NOT_APPLICABLE when no UK market
        if req.framework == Framework.UK_MDR and not uk_market_detected:
            keyword_results[req.req_id] = (
                ComplianceStatus.NOT_APPLICABLE,
                "No UK market detected — UK MDR requirements not applicable.",
            )
            continue

        # Get the section text for this requirement
        sec_text = ""
        if req.section in _SECTION_HEADING_PATTERNS:
            # Requirement maps to a PSUR section letter — find the matching key
            for sec_key, letters in _SECTION_KEY_TO_LETTERS.items():
                if req.section in letters and sec_key in section_texts:
                    sec_text = section_texts[sec_key]
                    break

        status, evidence = _keyword_prescreen_json(req, sec_text, full_text)
        keyword_results[req.req_id] = (status, evidence)

    # ── 2. LLM deep-analysis (batched per section) ────────────────────
    llm_results: Dict[str, Dict[str, str]] = {}

    if use_llm:
        # Group requirements by their section letter
        section_groups: Dict[str, List[Tuple[Requirement, ComplianceStatus, str]]] = {}
        for req in requirements:
            kw_status, kw_evidence = keyword_results[req.req_id]
            if kw_status == ComplianceStatus.NOT_APPLICABLE:
                continue
            section_groups.setdefault(req.section, []).append(
                (req, kw_status, kw_evidence)
            )

        for sec_letter, group in sorted(section_groups.items()):
            # Get the flattened text for this section
            sec_text = ""
            for sec_key, letters in _SECTION_KEY_TO_LETTERS.items():
                if sec_letter in letters and sec_key in section_texts:
                    sec_text = section_texts[sec_key]
                    break
            if not sec_text and sec_letter in _META_SECTION_LETTERS:
                sec_text = full_text

            prompt = _build_llm_json_audit_batch(group, sec_text)
            if prompt is None:
                continue

            needs_ids = [r.req_id for r, s, _ in group
                         if s in (ComplianceStatus.PARTIAL, ComplianceStatus.GAP)]

            if verbose:
                console.print(f"    [dim]LLM auditing section {sec_letter} "
                               f"({len(needs_ids)} reqs)...[/dim]")

            try:
                resp = create_message(
                    model=MODEL,
                    max_tokens=2000,
                    temperature=0.1,
                    system=_LLM_SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": prompt}],
                )
                parsed = _parse_llm_response(resp.content[0].text, needs_ids)
                llm_results.update(parsed)
            except Exception as e:
                if verbose:
                    console.print(f"    [yellow]LLM audit failed for section "
                                   f"{sec_letter}: {e}[/yellow]")

    # ── 3. Build findings ─────────────────────────────────────────────
    all_findings: List[AuditFinding] = []
    for req in requirements:
        kw_status, kw_evidence = keyword_results[req.req_id]

        if req.req_id in llm_results:
            llm_data = llm_results[req.req_id]
            try:
                final_status = ComplianceStatus(llm_data["status"])
            except (ValueError, KeyError):
                final_status = kw_status
            evidence = llm_data.get("evidence", kw_evidence)
            recommendation = llm_data.get("recommendation", "")
            llm_assessed = True
        else:
            final_status = kw_status
            evidence = kw_evidence
            recommendation = "" if final_status == ComplianceStatus.COMPLIANT else (
                f"Review {req.guidance_ref} for compliance guidance."
            )
            llm_assessed = False

        all_findings.append(AuditFinding(
            req_id=req.req_id,
            title=req.title,
            status=final_status,
            evidence=evidence[:300],
            recommendation=recommendation[:300],
            section=req.section,
            framework=req.framework.value,
            criticality=req.criticality.value,
            llm_assessed=llm_assessed,
        ))

    # ── 4. Build per-section audit results ────────────────────────────
    section_results: List[SectionAuditResult] = []
    for sec_key, letters in _SECTION_KEY_TO_LETTERS.items():
        sec_findings = [f for f in all_findings if f.section in letters]
        # Also include meta-section findings for cross-cutting concerns
        # (assign to M_findings_and_conclusions for remediation)
        if sec_key == "M_findings_and_conclusions":
            meta_findings = [f for f in all_findings if f.section in _META_SECTION_LETTERS]
            sec_findings.extend(meta_findings)

        result = SectionAuditResult(
            section_key=sec_key,
            findings=sec_findings,
            has_gaps=any(f.status == ComplianceStatus.GAP for f in sec_findings),
            has_critical_gaps=any(
                f.status == ComplianceStatus.GAP and f.criticality == "CRITICAL"
                for f in sec_findings
            ),
        )
        if result.has_gaps or any(f.status == ComplianceStatus.PARTIAL for f in sec_findings):
            result.build_remediation_prompt()
        section_results.append(result)

    # ── 5. Build full report ──────────────────────────────────────────
    report = AuditReport(
        psur_path="(in-memory)",
        audit_timestamp=time.strftime("%Y-%m-%dT%H:%M:%S"),
        uk_mdr_enabled=True,
        llm_enabled=use_llm,
        findings=all_findings,
    )
    report.compute_score()

    elapsed = time.time() - t0
    if verbose:
        console.print(
            f"    Score: [{'green' if report.compliance_score >= 80 else 'yellow' if report.compliance_score >= 60 else 'red'}]"
            f"{report.compliance_score}%[/] "
            f"({report.compliant}C / {report.partial}P / {report.gap}G / "
            f"{report.not_applicable}NA)  [{elapsed:.1f}s]"
        )

    return section_results, report


# ═══════════════════════════════════════════════════════════════════════════
# Main DOCX audit engine
# ═══════════════════════════════════════════════════════════════════════════

def run_audit(
    psur_path: str | Path,
    *,
    include_uk: bool = False,
    use_llm: bool = True,
    verbose: bool = False,
) -> AuditReport:
    """Run a full PSUR compliance audit.

    Args:
        psur_path: Path to a PSUR .docx file.
        include_uk: Include UK MDR requirements.
        use_llm: Use LLM for deep analysis (default True). Set False for keyword-only mode.
        verbose: Print progress to console.

    Returns:
        AuditReport with all findings, scores, and optional LLM summary.
    """
    psur_path = Path(psur_path)
    if not psur_path.exists():
        raise FileNotFoundError(f"PSUR file not found: {psur_path}")

    t0 = time.time()

    if verbose:
        console.print(f"\n[bold blue]{'='*56}[/bold blue]")
        console.print(f"[bold blue]  PSUR Auditor — MDCG 2022-21 Compliance Check[/bold blue]")
        console.print(f"[bold blue]{'='*56}[/bold blue]")
        console.print(f"\n  Document: {psur_path.name}")
        console.print(f"  UK MDR:   {'Enabled' if include_uk else 'Disabled'}")
        console.print(f"  LLM:      {'Enabled' if use_llm else 'Disabled (keyword only)'}\n")

    # ── 1. Parse the PSUR ─────────────────────────────────────────────
    if verbose:
        console.print("[bold]Parsing PSUR document...[/bold]")
    parser = PSURParser(psur_path)

    if verbose:
        console.print(f"  Words: {parser.get_word_count():,}")
        console.print(f"  Sections identified: {parser.get_section_count()}")
        console.print(f"  Tables: {parser.get_table_count()}")
        console.print(f"  Headings: {len(parser.headings)}\n")

    # ── 2. Build requirements checklist ───────────────────────────────
    requirements = build_requirements_checklist(include_uk=include_uk)
    if verbose:
        console.print(f"[bold]Auditing {len(requirements)} requirements...[/bold]\n")

    # ── 3. Keyword pre-screen ─────────────────────────────────────────
    keyword_results: Dict[str, Tuple[ComplianceStatus, str]] = {}
    for req in requirements:
        status, evidence = _keyword_prescreen(req, parser)
        keyword_results[req.req_id] = (status, evidence)

    if verbose:
        kw_comp = sum(1 for s, _ in keyword_results.values() if s == ComplianceStatus.COMPLIANT)
        kw_part = sum(1 for s, _ in keyword_results.values() if s == ComplianceStatus.PARTIAL)
        kw_gap = sum(1 for s, _ in keyword_results.values() if s == ComplianceStatus.GAP)
        console.print(f"  Keyword pre-screen: {kw_comp} compliant, {kw_part} partial, {kw_gap} gap")

    # ── 4. LLM deep-analysis (batched per section) ────────────────────
    llm_results: Dict[str, Dict[str, str]] = {}

    if use_llm:
        # Group requirements by section
        section_groups: Dict[str, List[Tuple[Requirement, ComplianceStatus, str]]] = {}
        for req in requirements:
            sec = req.section
            kw_status, kw_evidence = keyword_results[req.req_id]
            section_groups.setdefault(sec, []).append((req, kw_status, kw_evidence))

        # Count how many need LLM
        total_llm = sum(
            1 for req in requirements
            if keyword_results[req.req_id][0] in (ComplianceStatus.PARTIAL, ComplianceStatus.GAP)
        )
        if verbose:
            console.print(f"  LLM deep-analysis: {total_llm} requirements need assessment\n")

        for sec_key, group_reqs in sorted(section_groups.items()):
            prompt = _build_llm_audit_batch(group_reqs, parser, sec_key)
            if prompt is None:
                continue

            needs_llm_ids = [r.req_id for r, s, _ in group_reqs
                            if s in (ComplianceStatus.PARTIAL, ComplianceStatus.GAP)]

            if verbose:
                console.print(f"  [dim]LLM analysing section {sec_key} ({len(needs_llm_ids)} reqs)...[/dim]")

            try:
                resp = create_message(
                    model=MODEL,
                    max_tokens=2000,
                    temperature=0.1,
                    system=_LLM_SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": prompt}],
                )
                parsed = _parse_llm_response(resp.content[0].text, needs_llm_ids)
                llm_results.update(parsed)

                # Track token usage
                if "input" not in llm_results:
                    llm_results["_tokens"] = {"input": 0, "output": 0}
                # (stored in report below)
            except Exception as e:
                if verbose:
                    console.print(f"    [yellow]LLM call failed for section {sec_key}: {e}[/yellow]")

    # ── 5. Merge results and build findings ───────────────────────────
    report = AuditReport(
        psur_path=str(psur_path),
        audit_timestamp=time.strftime("%Y-%m-%dT%H:%M:%S"),
        uk_mdr_enabled=include_uk,
        llm_enabled=use_llm,
    )

    for req in requirements:
        kw_status, kw_evidence = keyword_results[req.req_id]

        # If LLM provided an override, use it
        if req.req_id in llm_results:
            llm_data = llm_results[req.req_id]
            try:
                final_status = ComplianceStatus(llm_data["status"])
            except (ValueError, KeyError):
                final_status = kw_status
            evidence = llm_data.get("evidence", kw_evidence)
            recommendation = llm_data.get("recommendation", "")
            llm_assessed = True
        else:
            final_status = kw_status
            evidence = kw_evidence
            recommendation = "" if final_status == ComplianceStatus.COMPLIANT else (
                f"Review {req.guidance_ref} for compliance guidance."
            )
            llm_assessed = False

        report.findings.append(AuditFinding(
            req_id=req.req_id,
            title=req.title,
            status=final_status,
            evidence=evidence[:300],
            recommendation=recommendation[:300],
            section=req.section,
            framework=req.framework.value,
            criticality=req.criticality.value,
            llm_assessed=llm_assessed,
        ))

    report.compute_score()

    # ── 6. LLM executive summary ──────────────────────────────────────
    if use_llm:
        if verbose:
            console.print("\n  [dim]Generating audit summary...[/dim]")
        report.llm_summary = _generate_audit_summary(report, parser)

    elapsed = time.time() - t0
    if verbose:
        console.print(f"\n  [dim]Audit completed in {elapsed:.1f}s[/dim]\n")

    return report


# ═══════════════════════════════════════════════════════════════════════════
# Console report rendering
# ═══════════════════════════════════════════════════════════════════════════

_STATUS_COLORS = {
    ComplianceStatus.COMPLIANT: "green",
    ComplianceStatus.PARTIAL: "yellow",
    ComplianceStatus.GAP: "red",
    ComplianceStatus.NOT_APPLICABLE: "dim",
}

_CRITICALITY_COLORS = {
    "CRITICAL": "bold red",
    "MAJOR": "bold yellow",
    "MINOR": "dim",
}


def print_audit_report(report: AuditReport) -> None:
    """Pretty-print the audit report to the console."""
    # ── Score card ────────────────────────────────────────────────────
    score_color = "green" if report.compliance_score >= 80 else (
        "yellow" if report.compliance_score >= 60 else "red"
    )
    console.print(Panel(
        f"[{score_color} bold]{report.compliance_score}% Compliance Score[/{score_color} bold]\n\n"
        f"  [green]{report.compliant}[/green] Compliant  "
        f"  [yellow]{report.partial}[/yellow] Partial  "
        f"  [red]{report.gap}[/red] Gap  "
        f"  [dim]{report.not_applicable}[/dim] N/A  "
        f"  (Total: {report.total_requirements})",
        title="PSUR Audit Results",
        border_style=score_color,
    ))

    # ── LLM Summary ──────────────────────────────────────────────────
    if report.llm_summary:
        console.print(Panel(
            report.llm_summary,
            title="Audit Executive Summary",
            border_style="blue",
        ))

    # ── Findings table ────────────────────────────────────────────────
    table = Table(title="Detailed Findings", show_lines=True)
    table.add_column("ID", style="bold", width=6)
    table.add_column("Title", width=30)
    table.add_column("Status", width=12)
    table.add_column("Crit.", width=8)
    table.add_column("Evidence", width=50)
    table.add_column("Recommendation", width=40)

    # Sort: GAP first, then PARTIAL, then COMPLIANT
    sort_order = {
        ComplianceStatus.GAP: 0,
        ComplianceStatus.PARTIAL: 1,
        ComplianceStatus.COMPLIANT: 2,
        ComplianceStatus.NOT_APPLICABLE: 3,
    }
    sorted_findings = sorted(report.findings, key=lambda f: sort_order.get(f.status, 9))

    for f in sorted_findings:
        status_color = _STATUS_COLORS.get(f.status, "white")
        crit_color = _CRITICALITY_COLORS.get(f.criticality, "white")
        llm_marker = " *" if f.llm_assessed else ""

        table.add_row(
            f.req_id,
            f.title,
            f"[{status_color}]{f.status.value}{llm_marker}[/{status_color}]",
            f"[{crit_color}]{f.criticality}[/{crit_color}]",
            f.evidence[:80] + ("..." if len(f.evidence) > 80 else ""),
            f.recommendation[:60] + ("..." if len(f.recommendation) > 60 else ""),
        )

    console.print(table)
    console.print("[dim]* = LLM-assessed finding[/dim]\n")

    # ── Token usage ───────────────────────────────────────────────────
    if report.llm_enabled and (report.token_usage.get("input", 0) > 0):
        console.print(
            f"[dim]LLM tokens: {report.token_usage['input']:,} in / "
            f"{report.token_usage['output']:,} out[/dim]"
        )


# ═══════════════════════════════════════════════════════════════════════════
# Standalone CLI
# ═══════════════════════════════════════════════════════════════════════════

def main() -> None:
    """Standalone CLI entry point (also usable via main.py audit)."""
    import argparse

    ap = argparse.ArgumentParser(
        description="LLM-powered PSUR Auditor — MDCG 2022-21 compliance check"
    )
    ap.add_argument("psur_docx", help="Path to PSUR .docx file")
    ap.add_argument("--uk-mdr", action="store_true",
                    help="Include UK MDR requirements")
    ap.add_argument("--no-llm", action="store_true",
                    help="Keyword-only mode (no LLM calls)")
    ap.add_argument("--output", "-o", type=str, default=None,
                    help="Save JSON report to this path")
    ap.add_argument("--verbose", "-v", action="store_true", default=True,
                    help="Verbose console output (default: on)")
    ap.add_argument("--ollama-model", type=str, default=None,
                    help="Use a local Ollama model instead of Claude")
    ap.add_argument("--ollama-url", type=str, default=None,
                    help="Ollama API base URL")
    args = ap.parse_args()

    # Ollama override
    if args.ollama_model:
        from llm_client import set_ollama_override
        set_ollama_override(args.ollama_model, url=args.ollama_url)

    report = run_audit(
        args.psur_docx,
        include_uk=args.uk_mdr,
        use_llm=not args.no_llm,
        verbose=args.verbose,
    )

    print_audit_report(report)

    if args.output:
        out_path = Path(args.output)
        with open(out_path, "w") as f:
            json.dump(asdict(report), f, indent=2, default=str)
        console.print(f"\n[green]Report saved: {out_path}[/green]")

    # Exit code: 0 if >=80%, 1 otherwise
    sys.exit(0 if report.compliance_score >= 80 else 1)


if __name__ == "__main__":
    main()
