"""TableMixin — all table fill/manipulation logic for the DOCX renderer."""
import copy
import logging
from typing import Any, Dict

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from imdrf_coder import strip_imdrf_code
from rendering._helpers import stringify, PH_RE

logger = logging.getLogger(__name__)


class TableMixin:
    """Provides all table-related fill methods for PSURTemplateRenderer."""

    # ==================================================================
    # Master dispatcher
    # ==================================================================
    def _fill_all_tables(self, psur: Dict[str, Any], values: Dict[str, str]):
        """Fill data into all tables by detecting content, not index."""
        sections = psur.get("sections", {})
        stats = psur.get("_statistics", {}) or {}

        sec_c = sections.get("C_volume_of_sales_and_population_exposure", {})
        t1_data = sec_c.get("table_1_sales_by_region", {})
        t1_fmt = t1_data.get("annual_format", t1_data) if isinstance(t1_data, dict) else {}
        sales_rows = t1_fmt.get("rows", []) if isinstance(t1_fmt, dict) else []
        date_ranges = t1_fmt.get("date_ranges", []) if isinstance(t1_fmt, dict) else []

        # ── Deterministic override: rebuild sales rows from _statistics ──
        # The LLM may emit a partial set of regions; the template has 11
        # fixed slots that must all reconcile to Worldwide. Use the pre-
        # computed section_c_region_rows when available so totals always
        # add up and no slot is silently em-dashed.
        det_rows = stats.get("section_c_region_rows") or []
        if det_rows:
            total_units = stats.get("total_units_sold", 0) or 0
            sales_rows = [
                {
                    "region": r.get("region", ""),
                    # Pull the 3 historical 12-month units directly from
                    # statistics so the template's "Preceding 12-Month"
                    # columns are filled deterministically when DB history
                    # is available.
                    "preceding_12_month_periods": [
                        r.get("units_p1"),
                        r.get("units_p2"),
                        r.get("units_p3"),
                    ],
                    "current_data_collection_period": r.get("units", 0),
                    "percent_of_global_sales": (
                        # Use pre-calculated pct_current if available (from skill),
                        # else calculate from units
                        r.get("pct_current") or (
                            round((r.get("units", 0) / total_units) * 100, 1)
                            if total_units > 0 else 0.0
                        )
                    ),
                }
                for r in det_rows
            ]
            # Surface period header labels for downstream use (renderer can
            # optionally stamp them into the Table 1 column headers).
            date_ranges = stats.get("section_c_period_labels") or date_ranges

        sec_d = sections.get("D_information_on_serious_incidents", {})
        sec_e = sections.get("E_customer_feedback", {})
        sec_f = sections.get("F_product_complaint_types_counts_and_rates", {})
        t7_data = sec_f.get("table_7_complaint_rate_and_count", {})
        t7_fmt = t7_data.get("annual_format", t7_data) if isinstance(t7_data, dict) else {}
        complaint_rows = t7_fmt.get("rows", []) if isinstance(t7_fmt, dict) else []
        grand_total = t7_fmt.get("grand_total", {}) if isinstance(t7_fmt, dict) else {}

        # ── Deterministic override: rebuild Table 7 rows + grand total ──
        # The LLM-built rows can drop uncoded complaints, breaking the
        # Grand Total = total_complaints reconciliation expected by auditors.
        # Use the pre-computed table7_rows when available.
        det_t7 = stats.get("table7_rows") or []
        if det_t7:
            complaint_rows = [
                {
                    "harm": r.get("harm", "No Harm"),
                    "medical_device_problem": r.get("medical_device_problem", ""),
                    "current_12_month_complaint_count": r.get("complaint_count", 0),
                    "current_12_month_complaint_rate": r.get("complaint_percentage", 0.0),
                    "max_expected_rate_of_occurrence_from_ract": (
                        r.get("ract_max_expected_rate")
                        if r.get("ract_max_expected_rate") is not None
                        else None
                    ),
                }
                for r in det_t7
            ]
            t7_total_count = sum(r.get("complaint_count", 0) for r in det_t7)
            wu = stats.get("total_units_sold", 0) or 0
            grand_total = {
                "complaint_count": t7_total_count,
                "complaint_rate": (
                    round((t7_total_count / wu) * 100, 4) if wu > 0 else 0.0
                ),
            }


        sec_h = sections.get("H_information_from_fsca", {})
        sec_i = sections.get("I_corrective_and_preventive_actions", {})
        sec_k = sections.get("K_review_of_external_databases_and_registries", {})
        sec_l = sections.get("L_pmcf", {})

        filled_tables: set = set()
        sales_table_filled = False
        complaint_table_filled = False

        for i, table in enumerate(self.doc.tables):
            header_text = self._get_table_header_text(table)
            ht = header_text.lower()
            logger.info(f"Table {i}: header='{header_text[:120]}'")

            if i in filled_tables:
                self._fill_table_placeholders(table, values)
                continue

            # ── IMDRF tables FIRST ──
            if "health impact" in ht or "annex f" in ht:
                logger.info(f"  → Table {i} matched: HEALTH_IMPACT")
                rows = sec_d.get("table_4_health_impact_by_investigation_conclusion", [])
                self._fill_health_impact_table(table, rows)
                filled_tables.add(i)

            elif "imdrf" in ht and ("annex a" in ht or "problem code" in ht):
                logger.info(f"  → Table {i} matched: INCIDENT_ANNEX_A")
                rows = sec_d.get("table_2_serious_incidents_by_imdrf_annex_a_by_region", [])
                self._fill_incident_table(table, rows)
                filled_tables.add(i)

            elif "imdrf" in ht and ("annex c" in ht or "cause code" in ht or "cause" in ht):
                logger.info(f"  → Table {i} matched: INCIDENT_ANNEX_C")
                rows = sec_d.get("table_3_serious_incidents_by_imdrf_annex_c_investigation_findings_by_region", [])
                self._fill_incident_table(table, rows)
                filled_tables.add(i)

            # ── Sales tables ──
            elif ("region" in ht and "imdrf" not in ht
                    and ("sales" in ht or "units" in ht or "volume" in ht
                         or "devices sold" in ht or "distributed" in ht
                         or "percent of global" in ht)):
                logger.info(f"  → Table {i} matched: SALES (filled={sales_table_filled})")
                if not sales_table_filled:
                    self._fill_sales_table(table, sales_rows, date_ranges, values)
                    sales_table_filled = True
                else:
                    self._clear_duplicate_table(table)
                filled_tables.add(i)

            elif "feedback" in ht and ("type" in ht or "source" in ht or "customer" in ht):
                logger.info(f"  → Table {i} matched: CUSTOMER_FEEDBACK")
                rows = sec_e.get("table_6_feedback_by_type_and_source", [])
                self._fill_generic_table(table, rows, start_row=1)
                filled_tables.add(i)

            elif (("harm" in ht and "medical device problem" in ht)
                    or ("complaint" in ht and ("rate" in ht or "count" in ht
                        or "harm" in ht or "mdp" in ht))):
                logger.info(f"  → Table {i} matched: COMPLAINT_RATE (filled={complaint_table_filled})")
                if not complaint_table_filled:
                    self._fill_complaint_table(table, complaint_rows, grand_total, values)
                    complaint_table_filled = True
                else:
                    self._clear_duplicate_table(table)
                filled_tables.add(i)

            elif "fsca" in ht or "field safety" in ht:
                logger.info(f"  → Table {i} matched: FSCA")
                rows = sec_h.get("table_8_fsca_initiated_current_period_and_open_fscas", [])
                self._fill_generic_table(table, rows, start_row=1)
                filled_tables.add(i)

            elif "capa" in ht or ("corrective" in ht and "preventive" in ht):
                logger.info(f"  → Table {i} matched: CAPA")
                rows = sec_i.get("table_9_capa_initiated_current_reporting_period", [])
                self._fill_generic_table(table, rows, start_row=1)
                filled_tables.add(i)

            elif (("database" in ht and ("registr" in ht or "total matches" in ht))
                    or "maude" in ht or "bfarm" in ht or "mhra" in ht):
                logger.info(f"  → Table {i} matched: EXTERNAL_DB")
                rows = sec_k.get("table_10_adverse_events_and_recalls", [])
                self._fill_generic_table(table, rows, start_row=1)
                filled_tables.add(i)

            elif "pmcf" in ht or "post-market clinical" in ht or "post market clinical" in ht:
                logger.info(f"  → Table {i} matched: PMCF")
                rows = sec_l.get("table_11_pmcf_activities", [])
                self._fill_generic_table(table, rows, start_row=1)
                filled_tables.add(i)

            else:
                logger.info(f"  → Table {i}: UNMATCHED")

            self._fill_table_placeholders(table, values)

    # ==================================================================
    # Standalone-tables splice
    # ==================================================================
    def _classify_table_header(self, header_text: str):
        """Categorise a table by its header text. Returns a category key
        used to match standalone-tables docx tables against template tables."""
        ht = (header_text or "").lower()
        if "health impact" in ht or "annex f" in ht:
            return "health_impact"
        if "imdrf" in ht and ("annex a" in ht or "problem code" in ht):
            return "incident_annex_a"
        if "imdrf" in ht and ("annex c" in ht or "cause code" in ht or "cause" in ht):
            return "incident_annex_c"
        if ("region" in ht and "imdrf" not in ht
                and ("sales" in ht or "units" in ht or "volume" in ht
                     or "devices sold" in ht or "distributed" in ht
                     or "percent of global" in ht)):
            return "sales"
        if "feedback" in ht and ("type" in ht or "source" in ht or "customer" in ht):
            return "customer_feedback"
        if (("harm" in ht and "medical device problem" in ht)
                or ("complaint" in ht and ("rate" in ht or "count" in ht
                    or "harm" in ht or "mdp" in ht))):
            return "complaint_rate"
        if "fsca" in ht or "field safety" in ht:
            return "fsca"
        if "capa" in ht or ("corrective" in ht and "preventive" in ht):
            return "capa"
        if (("database" in ht and ("registr" in ht or "total matches" in ht))
                or "maude" in ht or "bfarm" in ht or "mhra" in ht):
            return "external_db"
        if "pmcf" in ht or "post-market clinical" in ht or "post market clinical" in ht:
            return "pmcf"
        return None

    def _replace_tables_from_docx(self, tables_docx_path) -> int:
        """Replace data tables in self.doc with the equivalent tables from a
        standalone tables-only DOCX produced by build_tables_standalone.

        Matching is by header-text category (sales / IMDRF / complaint /
        FSCA / CAPA / external_db / PMCF / health_impact / customer_feedback).
        Returns the number of tables swapped.
        """
        from docx import Document as _Document
        from pathlib import Path as _Path

        path = _Path(tables_docx_path)
        if not path.exists():
            logger.warning("Standalone tables docx not found: %s", path)
            return 0

        src_doc = _Document(str(path))
        src_by_cat: Dict[str, list] = {}
        for stbl in src_doc.tables:
            cat = self._classify_table_header(self._get_table_header_text(stbl))
            if cat:
                src_by_cat.setdefault(cat, []).append(stbl._tbl)

        if not src_by_cat:
            logger.warning("No classifiable tables found in %s", path)
            return 0

        swapped = 0
        for dtbl in list(self.doc.tables):
            cat = self._classify_table_header(self._get_table_header_text(dtbl))
            if not cat or not src_by_cat.get(cat):
                continue
            src_el = src_by_cat[cat].pop(0)
            new_el = copy.deepcopy(src_el)
            parent = dtbl._tbl.getparent()
            if parent is None:
                continue
            parent.replace(dtbl._tbl, new_el)
            swapped += 1
            logger.info("Spliced standalone table into '%s' slot", cat)

        logger.info("Replaced %d tables from %s", swapped, path.name)
        return swapped

    # ==================================================================
    # Header detection
    # ==================================================================
    def _get_table_header_text(self, table) -> str:
        """Get text from the first 2 rows of a table for identification."""
        tbl_el = table._tbl
        trs = list(tbl_el.iterchildren(qn("w:tr")))
        texts = []
        for tr in trs[:2]:
            texts.append(self._get_row_text(tr))
        return " ".join(texts)

    # ==================================================================
    # Placeholder substitution in table cells
    # ==================================================================
    def _fill_table_placeholders(self, table, values: Dict[str, str]):
        """Replace {{PLACEHOLDER}} markers in all table cells via raw XML."""
        tbl_el = table._tbl
        for tr in tbl_el.iterchildren(qn("w:tr")):
            for tc in tr.iterchildren(qn("w:tc")):
                for p in tc.iterchildren(qn("w:p")):
                    runs = p.findall(f".//{qn('w:r')}")
                    if not runs:
                        continue
                    texts = []
                    for r in runs:
                        t_el = r.find(qn("w:t"))
                        if t_el is not None and t_el.text:
                            texts.append(t_el.text)
                    full = "".join(texts)
                    if "{{" not in full:
                        continue
                    new_full = PH_RE.sub(
                        lambda m: values.get(m.group(1), m.group(0)), full)
                    if new_full == full:
                        continue
                    first = True
                    for r in runs:
                        t_el = r.find(qn("w:t"))
                        if t_el is not None:
                            if first:
                                t_el.text = new_full
                                t_el.set(qn("xml:space"), "preserve")
                                first = False
                            else:
                                t_el.text = ""

    # ==================================================================
    # Duplicate table clearing
    # ==================================================================
    def _clear_duplicate_table(self, table):
        """Remove a duplicate table entirely from the document.

        When the template contains both annual and biennial variants, the
        second match is the unused variant and should be deleted completely
        rather than filled with dashes.
        """
        tbl_el = table._tbl
        parent = tbl_el.getparent()
        if parent is not None:
            parent.remove(tbl_el)
            logger.info("Removed duplicate table variant from document")

    # ==================================================================
    # Sales table
    # ==================================================================
    def _fill_sales_table(self, table, sales_rows: list,
                          date_ranges: list, values: Dict[str, str]):
        """Fill the sales-by-region table."""
        tbl_el = table._tbl
        trs = list(tbl_el.iterchildren(qn("w:tr")))

        region_data = {}
        for row in sales_rows:
            if not isinstance(row, dict):
                continue
            region = row.get("region", "")
            preceding = row.get("preceding_12_month_periods", [])
            if not isinstance(preceding, list):
                preceding = []
            current = row.get("current_data_collection_period", "")
            pct = row.get("percent_of_global_sales", "")

            def _fmt_units(v):
                """Match the 'current' column's number formatting for trend cells."""
                if v is None or v == "":
                    return "Data not available"
                if isinstance(v, (int, float)):
                    return f"{int(v):,}" if float(v) >= 0 else stringify(v)
                return stringify(v)

            if isinstance(current, (int, float)):
                current_text = f"{int(current):,}"
            elif current in (None, ""):
                current_text = "0"
            else:
                current_text = stringify(current)

            if isinstance(pct, (int, float)):
                pct_text = f"{float(pct):.1f}%"
            elif pct in (None, ""):
                pct_text = "0.0%"
            else:
                pct_text = stringify(pct)
                if not pct_text.endswith("%"):
                    pct_text = f"{pct_text}%"

            region_data[region] = {
                "p1": _fmt_units(preceding[0]) if len(preceding) > 0 else "Data not available",
                "p2": _fmt_units(preceding[1]) if len(preceding) > 1 else "Data not available",
                "p3": _fmt_units(preceding[2]) if len(preceding) > 2 else "Data not available",
                "current": current_text,
                "pct": pct_text,
            }

        template_regions = [
            "EEA+TR+XI", "Australia", "Brazil", "Canada", "China",
            "Japan", "UK", "United States",
            "Unknown / Unattributed",
            "Rest of World", "Worldwide",
        ]

        for ri, region_name in enumerate(template_regions):
            data_row_idx = ri + 3
            if data_row_idx >= len(trs):
                break

            tr = trs[data_row_idx]
            tcs = list(tr.iterchildren(qn("w:tc")))

            if region_name is None:
                self._set_cell_text(tcs[0], "N/A")
                for tc in tcs[1:]:
                    self._set_cell_text(tc, "N/A")
                continue

            data = region_data.get(region_name)
            if not data:
                for rk, rv in region_data.items():
                    if region_name.lower() in rk.lower() or rk.lower() in region_name.lower():
                        data = rv
                        break

            if data and len(tcs) >= 6:
                # Stamp the region label so repurposed slots (e.g. the blank
                # divider row now used for "Unknown / Unattributed") get a
                # visible name. Existing rows already have correct labels in
                # the template; overwriting with the same string is a no-op.
                self._set_cell_text(tcs[0], region_name)
                self._set_cell_text(tcs[1], data["p1"])
                self._set_cell_text(tcs[2], data["p2"])
                self._set_cell_text(tcs[3], data["p3"])
                self._set_cell_text(tcs[4], data["current"])
                self._set_cell_text(tcs[5], data["pct"])
            elif len(tcs) >= 6:
                self._set_cell_text(tcs[0], region_name)
                self._set_cell_text(tcs[1], "Data not available")
                self._set_cell_text(tcs[2], "Data not available")
                self._set_cell_text(tcs[3], "Data not available")
                self._set_cell_text(tcs[4], "0")
                self._set_cell_text(tcs[5], "0.0%")

        # Ensure the date-label row contains no blank cells after template cleanup.
        if len(trs) > 2:
            tcs = list(trs[2].iterchildren(qn("w:tc")))
            labels = list(date_ranges or [])
            if len(tcs) >= 6:
                self._set_cell_text(tcs[0], "Region")
                self._set_cell_text(tcs[1], stringify(labels[0]) if len(labels) > 0 else "Data not available")
                self._set_cell_text(tcs[2], stringify(labels[1]) if len(labels) > 1 else "Data not available")
                self._set_cell_text(tcs[3], stringify(labels[2]) if len(labels) > 2 else "Data not available")
                self._set_cell_text(tcs[4], stringify(labels[-1]) if labels else "Current data collection period")
                self._set_cell_text(tcs[5], "12-Month Percent of Global Sales")
            if len(table.rows) > 2:
                cells = table.rows[2].cells
                cell_values = [
                    "Region",
                    stringify(labels[0]) if len(labels) > 0 else "Data not available",
                    stringify(labels[1]) if len(labels) > 1 else "Data not available",
                    stringify(labels[2]) if len(labels) > 2 else "Data not available",
                    stringify(labels[-1]) if labels else "Current data collection period",
                    "12-Month Percent of Global Sales",
                ]
                for ci, cell in enumerate(cells[:len(cell_values)]):
                    cell.text = cell_values[ci]
            if len(table.rows) > 1:
                cells = table.rows[1].cells
                label = stringify(labels[-1]) if labels else "Current data collection period"
                if len(cells) >= 5:
                    cells[4].text = label
                if len(cells) >= 6:
                    cells[5].text = "12-Month Percent of Global Sales"
            if len(table.rows) > 0:
                cells = table.rows[0].cells
                label = stringify(labels[-1]) if labels else "Current data collection period"
                if len(cells) >= 5:
                    cells[4].text = label
                if len(cells) >= 6:
                    cells[5].text = "12-Month Percent of Global Sales"
            trs = list(tbl_el.iterchildren(qn("w:tr")))
            if len(trs) > 2:
                tbl_el.remove(trs[2])
            trs = list(tbl_el.iterchildren(qn("w:tr")))
            if len(trs) > 1:
                tbl_el.remove(trs[1])

    # ==================================================================
    # Complaint table
    # ==================================================================
    def _fill_complaint_table(self, table, complaint_rows: list,
                               grand_total: dict, values: Dict[str, str]):
        """Fill the complaint rate table."""
        tbl_el = table._tbl
        trs = list(tbl_el.iterchildren(qn("w:tr")))

        if not complaint_rows:
            if len(trs) > 2:
                tcs = list(trs[2].iterchildren(qn("w:tc")))
                if tcs:
                    self._set_cell_text(tcs[0], "None reported during this period.")
            return

        # Deduplicate MDP rows: same harm + same MDP should only appear once
        # (take the row with the highest complaint count if duplicated)
        seen_keys = {}
        deduped_rows = []
        for row in complaint_rows:
            if not isinstance(row, dict):
                continue
            harm = row.get("harm", "No Health Consequence or Impact")
            mdp = row.get("medical_device_problem", "")
            key = f"{harm}|{mdp}"
            if key in seen_keys:
                # Keep the row with higher count
                existing = seen_keys[key]
                if row.get("current_12_month_complaint_count", 0) > existing.get("current_12_month_complaint_count", 0):
                    deduped_rows.remove(existing)
                    deduped_rows.append(row)
                    seen_keys[key] = row
            else:
                seen_keys[key] = row
                deduped_rows.append(row)

        harm_groups: Dict[str, list] = {}
        for row in deduped_rows:
            harm = row.get("harm", "No Health Consequence or Impact")
            harm_groups.setdefault(harm, []).append(row)

        noharm_key = None
        other_harm_keys = []
        for hk in harm_groups:
            if "no health" in hk.lower() or "no harm" in hk.lower():
                noharm_key = hk
            else:
                other_harm_keys.append(hk)

        def _fmt_count_rate(count, rate):
            if isinstance(count, (int, float)):
                count_text = f"{int(count):,}"
            else:
                count_text = stringify(count or 0)
            if isinstance(rate, (int, float)):
                rate_text = f"{float(rate):.4f}%"
            elif rate in (None, ""):
                rate_text = "0.0000%"
            else:
                rate_text = stringify(rate)
                if not rate_text.endswith("%"):
                    rate_text = f"{rate_text}%"
            return f"{rate_text} ({count_text})"

        def _fmt_ract(row, is_header=False, is_total=False):
            if is_total or is_header:
                return "N/A"
            ract = row.get("max_expected_rate_of_occurrence_from_ract")
            if ract is None or ract == "" or ract == "N/A":
                return "N/A - RACT not provided"
            if isinstance(ract, (int, float)):
                return f"{float(ract):.4f}%"
            return stringify(ract)

        rendered_rows = []
        for harm_key in other_harm_keys + ([noharm_key] if noharm_key else []):
            if not harm_key:
                continue
            rows = harm_groups[harm_key]
            total_count = sum(r.get("current_12_month_complaint_count", 0) or 0 for r in rows)
            total_rate = sum(r.get("current_12_month_complaint_rate", 0.0) or 0.0 for r in rows)
            rendered_rows.append({
                "label": strip_imdrf_code(harm_key),
                "rate_count": _fmt_count_rate(total_count, total_rate),
                "ract": _fmt_ract({}, is_header=True),
                "kind": "header",
            })
            for mdp_row in rows:
                rendered_rows.append({
                    "label": "  " + strip_imdrf_code(mdp_row.get("medical_device_problem", "")),
                    "rate_count": _fmt_count_rate(
                        mdp_row.get("current_12_month_complaint_count", 0),
                        mdp_row.get("current_12_month_complaint_rate", 0.0),
                    ),
                    "ract": _fmt_ract(mdp_row),
                    "kind": "mdp",
                })

        rendered_rows.append({
            "label": "Grand Total",
            "rate_count": _fmt_count_rate(
                grand_total.get("complaint_count", 0),
                grand_total.get("complaint_rate", 0.0),
            ),
            "ract": _fmt_ract({}, is_total=True),
            "kind": "total",
        })

        template_tr = copy.deepcopy(trs[2] if len(trs) > 2 else trs[-1])
        self._trim_table_rows(tbl_el, keep_count=1)
        for row_data in rendered_rows:
            new_tr = copy.deepcopy(template_tr)
            tbl_el.append(new_tr)
            tcs = list(new_tr.iterchildren(qn("w:tc")))
            if len(tcs) >= 1:
                self._set_cell_text(tcs[0], row_data["label"])
            if len(tcs) >= 2:
                self._set_cell_text(tcs[1], row_data["rate_count"])
            if len(tcs) >= 3:
                self._set_cell_text(tcs[2], row_data["ract"])
            for tc in tcs[3:]:
                self._set_cell_text(tc, "N/A")
            if row_data["kind"] in {"header", "total"}:
                self._set_row_bold(new_tr)
                self._set_row_shading(new_tr, "F2F2F2")
        return

        def _fill_mdp_row(tr_el, mdp_row: dict):
            tcs = list(tr_el.iterchildren(qn("w:tc")))
            if len(tcs) < 3:
                return
            mdp_name = strip_imdrf_code(mdp_row.get("medical_device_problem", ""))
            self._set_cell_text(tcs[0], mdp_name)
            count = mdp_row.get("current_12_month_complaint_count", "")
            rate = mdp_row.get("current_12_month_complaint_rate", "")
            count_str = f"{count:,}" if isinstance(count, (int, float)) else stringify(count)
            rate_str = f"{rate}%" if rate else ""
            self._set_cell_text(tcs[1], f"{count_str} ({rate_str})" if rate_str else count_str)
            ract = mdp_row.get("max_expected_rate_of_occurrence_from_ract")
            if ract is not None and ract != "" and ract != "N/A":
                # Format RACT value with occurrence code if available
                oc_code = mdp_row.get("occurrence_code", "")
                oc_max = mdp_row.get("occurrence_max_expected_rate", "")
                if oc_max and oc_code:
                    ract_str = f"≤{oc_max} ({oc_code})"
                else:
                    ract_str = stringify(ract)
                self._set_cell_text(tcs[2], ract_str)
            else:
                self._set_cell_text(tcs[2], "N/A — RACT not provided")

        def _fill_harm_header_row(tr_el, harm_name: str, total_count=None):
            tcs = list(tr_el.iterchildren(qn("w:tc")))
            if tcs:
                self._set_cell_text(tcs[0], strip_imdrf_code(harm_name))
            if len(tcs) >= 2 and total_count is not None:
                self._set_cell_text(tcs[1], f"{total_count:,}" if isinstance(total_count, (int, float)) else stringify(total_count))

        # Harm A (rows 2-4)
        if other_harm_keys:
            harm_a_key = other_harm_keys[0]
            harm_a_rows = harm_groups[harm_a_key]
            harm_a_total = sum(r.get("current_12_month_complaint_count", 0) for r in harm_a_rows)
            if len(trs) > 2:
                _fill_harm_header_row(trs[2], harm_a_key, harm_a_total)
            if len(trs) > 3 and len(harm_a_rows) > 0:
                _fill_mdp_row(trs[3], harm_a_rows[0])
            if len(trs) > 4 and len(harm_a_rows) > 1:
                _fill_mdp_row(trs[4], harm_a_rows[1])
            if len(harm_a_rows) > 2:
                insert_after = trs[4]
                for extra in harm_a_rows[2:]:
                    new_tr = copy.deepcopy(trs[3])
                    insert_after.addnext(new_tr)
                    _fill_mdp_row(new_tr, extra)
                    insert_after = new_tr
                trs = list(tbl_el.iterchildren(qn("w:tr")))
        else:
            if len(trs) > 2:
                tcs = list(trs[2].iterchildren(qn("w:tc")))
                if tcs:
                    self._set_cell_text(tcs[0], "N/A")

        # Harm B
        harm_b_start = None
        for idx, tr in enumerate(trs):
            text = self._get_row_text(tr)
            if "T7_HARM_B" in text and "MDP" not in text:
                harm_b_start = idx
                break

        if harm_b_start and len(other_harm_keys) > 1:
            harm_b_key = other_harm_keys[1]
            harm_b_rows = harm_groups[harm_b_key]
            harm_b_total = sum(r.get("current_12_month_complaint_count", 0) for r in harm_b_rows)
            _fill_harm_header_row(trs[harm_b_start], harm_b_key, harm_b_total)
            if harm_b_start + 1 < len(trs) and len(harm_b_rows) > 0:
                _fill_mdp_row(trs[harm_b_start + 1], harm_b_rows[0])
            if harm_b_start + 2 < len(trs) and len(harm_b_rows) > 1:
                _fill_mdp_row(trs[harm_b_start + 2], harm_b_rows[1])
        elif harm_b_start:
            tcs = list(trs[harm_b_start].iterchildren(qn("w:tc")))
            if tcs:
                self._set_cell_text(tcs[0], "N/A")

        # No Health Consequence section
        noharm_start = None
        for idx, tr in enumerate(trs):
            text = self._get_row_text(tr)
            if "No Health Consequence" in text:
                noharm_start = idx
                break

        if noharm_start is not None and noharm_key:
            noharm_rows = harm_groups[noharm_key]
            noharm_total = sum(r.get("current_12_month_complaint_count", 0) for r in noharm_rows)
            _fill_harm_header_row(trs[noharm_start], "No Health Consequence or Impact", noharm_total)

            mdp_slot_idx = noharm_start + 1
            filled = 0
            while mdp_slot_idx < len(trs) and filled < min(2, len(noharm_rows)):
                text = self._get_row_text(trs[mdp_slot_idx])
                if "Grand Total" in text:
                    break
                _fill_mdp_row(trs[mdp_slot_idx], noharm_rows[filled])
                filled += 1
                mdp_slot_idx += 1

            if len(noharm_rows) > 2:
                insert_after_idx = noharm_start + 2
                if insert_after_idx < len(trs):
                    insert_after = trs[insert_after_idx]
                    template_mdp_tr = trs[noharm_start + 1]
                    for extra_row in noharm_rows[2:]:
                        new_tr = copy.deepcopy(template_mdp_tr)
                        insert_after.addnext(new_tr)
                        _fill_mdp_row(new_tr, extra_row)
                        insert_after = new_tr
                    trs = list(tbl_el.iterchildren(qn("w:tr")))

        # Grand Total row
        for tr in reversed(trs):
            text = self._get_row_text(tr)
            if "Grand Total" in text:
                tcs = list(tr.iterchildren(qn("w:tc")))
                total_count = grand_total.get("complaint_count", "")
                total_rate = grand_total.get("complaint_rate", "")
                if len(tcs) >= 2:
                    count_str = f"{total_count:,}" if isinstance(total_count, (int, float)) else stringify(total_count)
                    rate_str = f"{total_rate}%" if total_rate else ""
                    self._set_cell_text(tcs[1], f"{count_str} ({rate_str})" if rate_str else count_str)
                break

    # ==================================================================
    # Incident table
    # ==================================================================
    def _fill_incident_table(self, table, rows: list):
        """Fill serious incident tables. If no data, write 'None reported'."""
        tbl_el = table._tbl
        trs = list(tbl_el.iterchildren(qn("w:tr")))

        def _norm_region(value):
            return "".join(ch for ch in stringify(value).lower() if ch.isalnum())

        def _fmt_rate(value):
            if isinstance(value, (int, float)):
                return f"{float(value):.4f}%"
            if value in (None, ""):
                return "0.0000%"
            text = stringify(value)
            return text if text.endswith("%") else f"{text}%"

        if not rows:
            _region_keywords = (
                "EEA", "UK", "Worldwide", "United States", "US", "Rest of World",
                "Australia", "Canada", "Japan", "China", "Brazil",
            )
            for tr in trs[1:]:
                tcs = list(tr.iterchildren(qn("w:tc")))
                text = self._get_row_text(tr)
                if any(region in text for region in _region_keywords):
                    for ci, tc in enumerate(tcs[1:], start=1):
                        self._set_cell_text(tc, "0.0000%" if ci == 3 else "0")
                elif text.strip():
                    for ci, tc in enumerate(tcs[1:], start=1):
                        self._set_cell_text(tc, "0.0000%" if ci == 3 else "0")
                else:
                    if tcs:
                        self._set_cell_text(tcs[0], "None reported during this period.")
            return

        for row_data in rows:
            if not isinstance(row_data, dict):
                continue
            region = row_data.get("region", "")
            for tr in trs[1:]:
                text = self._get_row_text(tr)
                if region and _norm_region(region) in _norm_region(text):
                    tcs = list(tr.iterchildren(qn("w:tc")))
                    col_keys = ["imdrf_code_and_term", "count", "rate", "complaint_number"]
                    for ci, key in enumerate(col_keys):
                        if ci + 1 < len(tcs):
                            val = row_data.get(key, "")
                            if key == "rate":
                                val_str = _fmt_rate(val)
                            else:
                                val_str = stringify(val) if val not in (None, "") else "0"
                            if key == "imdrf_code_and_term":
                                val_str = strip_imdrf_code(val_str)
                            self._set_cell_text(tcs[ci + 1], val_str)
        self._trim_table_rows(tbl_el, keep_count=min(len(trs), len(rows) + 1))

    def _fill_health_impact_table(self, table, rows: list):
        """Fill the health-impact serious incident table from skill rows."""
        tbl_el = table._tbl
        trs = list(tbl_el.iterchildren(qn("w:tr")))
        if not rows:
            rows = [
                {"region": "EEA+TR+XI", "health_impact": "N/A", "count": 0, "rate": 0.0, "complaint_number": "N/A"},
                {"region": "UK", "health_impact": "N/A", "count": 0, "rate": 0.0, "complaint_number": "N/A"},
                {"region": "Worldwide", "health_impact": "N/A", "count": 0, "rate": 0.0, "complaint_number": "N/A"},
            ]

        template_tr = copy.deepcopy(trs[1] if len(trs) > 1 else trs[-1])
        self._trim_table_rows(tbl_el, keep_count=1)
        for row_data in rows:
            tr = copy.deepcopy(template_tr)
            tbl_el.append(tr)
            tcs = list(tr.iterchildren(qn("w:tc")))
            region = row_data.get("region") or "N/A"
            health_impact = strip_imdrf_code(row_data.get("health_impact") or "N/A")
            values = [
                f"{region} - {health_impact}",
                stringify(row_data.get("count", 0)),
                f"{float(row_data.get('rate', 0) or 0):.4f}%",
                row_data.get("complaint_number") or "N/A",
                "N/A",
                "N/A",
            ]
            for ci, tc in enumerate(tcs):
                self._set_cell_text(tc, stringify(values[ci]) if ci < len(values) else "N/A")

        for ri, row_data in enumerate(rows, start=1):
            if ri >= len(table.rows):
                break
            cells = table.rows[ri].cells
            region = row_data.get("region") or "N/A"
            health_impact = strip_imdrf_code(row_data.get("health_impact") or "N/A")
            cell_values = [
                f"{region} - {health_impact}",
                stringify(row_data.get("count", 0)),
                f"{float(row_data.get('rate', 0) or 0):.4f}%",
                row_data.get("complaint_number") or "N/A",
                "N/A",
                "N/A",
            ]
            for ci, cell in enumerate(cells[:len(cell_values)]):
                cell.text = stringify(cell_values[ci])

    # ==================================================================
    # Generic table fill
    # ==================================================================
    def _fill_generic_table(self, table, rows: list, start_row: int = 1):
        """Fill a simple table with dict rows."""
        tbl_el = table._tbl
        trs = list(tbl_el.iterchildren(qn("w:tr")))

        if not rows:
            if start_row < len(trs):
                tcs = list(trs[start_row].iterchildren(qn("w:tc")))
                if tcs:
                    self._set_cell_text(tcs[0], "None reported during this period.")
                    for tc in tcs[1:]:
                        self._set_cell_text(tc, "N/A")
                    self._trim_table_rows(tbl_el, keep_count=start_row + 1)
            return

        if trs:
            header_tcs = list(trs[0].iterchildren(qn("w:tc")))
            headers = [self._get_cell_text(tc).lower().replace(" ", "_").replace("/", "_")
                       for tc in header_tcs]
        else:
            headers = []

        for ri, row_data in enumerate(rows):
            if not isinstance(row_data, dict):
                continue
            data_row_idx = start_row + ri
            if data_row_idx >= len(trs):
                template_tr = trs[start_row] if start_row < len(trs) else trs[-1]
                new_tr = copy.deepcopy(template_tr)
                tbl_el.append(new_tr)
                trs = list(tbl_el.iterchildren(qn("w:tr")))

            tr = trs[data_row_idx]
            tcs = list(tr.iterchildren(qn("w:tc")))

            row_keys = list(row_data.keys())
            for ci, tc in enumerate(tcs):
                val = None
                if ci < len(headers):
                    for rk, rv in row_data.items():
                        rk_norm = rk.lower().replace(" ", "_").replace("/", "_")
                        if rk_norm == headers[ci] or headers[ci] in rk_norm or rk_norm in headers[ci]:
                            val = rv
                            break
                if val is None and ci < len(row_keys):
                    val = row_data[row_keys[ci]]

                if val is not None:
                    self._set_cell_text(tc, stringify(val))
                else:
                    self._set_cell_text(tc, "N/A")

        self._trim_table_rows(tbl_el, keep_count=start_row + len([r for r in rows if isinstance(r, dict)]))

    # ==================================================================
    # XML cell helpers
    # ==================================================================
    def _trim_table_rows(self, tbl_el, keep_count: int):
        """Remove all rows after keep_count so template debris cannot remain."""
        trs = list(tbl_el.iterchildren(qn("w:tr")))
        for tr in trs[keep_count:]:
            tbl_el.remove(tr)

    def _set_row_bold(self, tr):
        """Apply bold formatting to all existing runs in a row."""
        for r in tr.iter(qn("w:r")):
            r_pr = r.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                r.insert(0, r_pr)
            if r_pr.find(qn("w:b")) is None:
                r_pr.append(OxmlElement("w:b"))

    def _set_row_shading(self, tr, fill: str):
        """Apply a solid background fill to every cell in a row."""
        for tc in tr.iterchildren(qn("w:tc")):
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), fill)

    def _set_cell_text(self, tc, text: str):
        """Set the text content of a table cell (w:tc), preserving formatting."""
        paragraphs = list(tc.iterchildren(qn("w:p")))
        if not paragraphs:
            return

        p = paragraphs[0]
        runs = p.findall(f".//{qn('w:r')}")

        if runs:
            t_el = runs[0].find(qn("w:t"))
            if t_el is None:
                t_el = OxmlElement("w:t")
                runs[0].append(t_el)
            t_el.text = text
            t_el.set(qn("xml:space"), "preserve")
            for r in runs[1:]:
                t_el2 = r.find(qn("w:t"))
                if t_el2 is not None:
                    t_el2.text = ""
        else:
            r = OxmlElement("w:r")
            t_el = OxmlElement("w:t")
            t_el.text = text
            t_el.set(qn("xml:space"), "preserve")
            r.append(t_el)
            p.append(r)

    def _get_cell_text(self, tc) -> str:
        """Get text content of a w:tc element."""
        texts = []
        for t in tc.iter(qn("w:t")):
            if t.text:
                texts.append(t.text)
        return "".join(texts).strip()

    def _get_row_text(self, tr) -> str:
        """Get concatenated text of all cells in a w:tr element."""
        texts = []
        for tc in tr.iterchildren(qn("w:tc")):
            texts.append(self._get_cell_text(tc))
        return " ".join(texts)
