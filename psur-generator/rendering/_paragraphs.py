"""ParagraphMixin — paragraph/run replacement and chart insertion for DOCX renderer."""
import logging
from pathlib import Path
from typing import Any, Dict

from docx.shared import Inches
from docx.oxml import OxmlElement

from rendering._helpers import PH_RE

logger = logging.getLogger(__name__)


class ParagraphMixin:
    """Provides paragraph-level fill, chart insertion, and header/footer update."""

    def _fill_paragraph(self, para, values: Dict[str, str]):
        """Replace {{PLACEHOLDER}} markers and [INSERT CHART] directives in a paragraph's runs."""
        full_text = para.text

        # [INSERT CHART] text-based directives
        upper_text = full_text.upper()
        if "[INSERT CHART" in upper_text:
            chart_key = None
            if "SALES" in upper_text or "VOLUME" in upper_text:
                chart_key = "sales_trend"
            elif "TREND" in upper_text or "UCL" in upper_text or "COMPLAINT RATE" in upper_text:
                chart_key = "trend_ucl"
            if chart_key and chart_key in self.chart_paths:
                chart_path = self.chart_paths[chart_key]
                for run in para.runs:
                    run.text = ""
                if not self._should_insert_chart(chart_key):
                    para.text = self._chart_note(chart_key)
                elif Path(chart_path).exists():
                    para.text = self._chart_caption(chart_key)
                    self._insert_picture_after(para, Path(chart_path), self._chart_note(chart_key))
                    self._placed_chart_keys.add(chart_key)
                return

        if "{{" not in full_text:
            return

        # {{PLACEHOLDER}} style chart placeholders (legacy)
        chart_map = {
            "sales_trend": "C_SALES_CHART_PLACEHOLDER",
            "trend_ucl": "G_TREND_CHART_PLACEHOLDER",
        }
        for chart_key, chart_path in self.chart_paths.items():
            ph_name = chart_map.get(chart_key)
            if ph_name and f"{{{{{ph_name}}}}}" in full_text:
                for run in para.runs:
                    run.text = ""
                if not self._should_insert_chart(chart_key):
                    para.text = self._chart_note(chart_key)
                elif Path(chart_path).exists():
                    para.text = self._chart_caption(chart_key)
                    self._insert_picture_after(para, Path(chart_path), self._chart_note(chart_key))
                    self._placed_chart_keys.add(chart_key)
                return

        self._replace_in_runs(para, values)

    def _insert_unplaced_charts(self, values: Dict[str, str]):
        """Insert generated charts into their report sections.

        The FormQAR-054 template is inconsistent about chart placeholders, so
        the renderer must place generated PNGs deterministically by section
        anchor instead of relying only on literal {{...}} placeholders.
        """
        if not self.chart_paths:
            return

        chart_plan = [
            ("sales_trend", "Section C: Volume Of Sales and Population Exposure", "Figure C-1. Sales trend by reporting period."),
            ("complaints_region", "Section F: Product Complaint Types, Complaint Counts, and Complaint Rates", "Figure F-1. Complaints by region."),
            ("harm_distribution", "Section F: Product Complaint Types, Complaint Counts, and Complaint Rates", "Figure F-2. Complaint distribution by harm category."),
            ("top_mdps", "Section F: Product Complaint Types, Complaint Counts, and Complaint Rates", "Figure F-3. Top medical device problem categories."),
            ("trend_ucl", "Overall Monthly Complaint Rate Trending", "Figure G-1. Monthly complaint rate control chart."),
            ("rate_occurrence", "Overall Monthly Complaint Rate Trending", "Figure G-2. Complaint rate against occurrence reference bands."),
            ("harm_trend", "Overall Monthly Complaint Rate Trending", "Figure G-3. Harm-category trend by month."),
            ("per_period", "Overall Monthly Complaint Rate Trending", "Figure G-4. Complaint count and rate by reporting period."),
            ("ract_matrix", "Overall Monthly Complaint Rate Trending", "Figure G-5. RACT severity and occurrence matrix."),
        ]

        inserted_by_anchor: Dict[str, object] = {}
        for chart_key, anchor_text, caption in chart_plan:
            if chart_key in getattr(self, "_placed_chart_keys", set()):
                continue
            if not self._should_insert_chart(chart_key):
                if chart_key in {"complaints_region", "per_period", "ract_matrix"}:
                    anchor_para = inserted_by_anchor.get(anchor_text)
                    if anchor_para is None:
                        anchor_para = self._find_chart_anchor(anchor_text, values)
                    if anchor_para is not None:
                        inserted_by_anchor[anchor_text] = self._insert_chart_limitation_after(
                            anchor_para,
                            self._chart_note(chart_key),
                        )
                    self._placed_chart_keys.add(chart_key)
                continue
            chart_path = self.chart_paths.get(chart_key)
            if not chart_path or not Path(chart_path).exists():
                continue

            anchor_para = inserted_by_anchor.get(anchor_text)
            if anchor_para is None:
                anchor_para = self._find_chart_anchor(anchor_text, values)
            if anchor_para is None:
                logger.warning("Could not find insertion point for chart %s", chart_key)
                continue

            inserted_by_anchor[anchor_text] = self._insert_chart_after(
                anchor_para,
                Path(chart_path),
                self._chart_caption(chart_key),
            )
            self._placed_chart_keys.add(chart_key)
            logger.info("Inserted chart %s after '%s'", chart_key, anchor_text)

    def _chart_caption(self, chart_key: str) -> str:
        context = getattr(self, "_chart_context", {})
        if chart_key in context and context[chart_key].get("caption"):
            return context[chart_key]["caption"]
        return f"Figure. {chart_key.replace('_', ' ').title()}."

    def _chart_note(self, chart_key: str) -> str:
        context = getattr(self, "_chart_context", {})
        return context.get(chart_key, {}).get("note", "")

    def _should_insert_chart(self, chart_key: str) -> bool:
        context = getattr(self, "_chart_context", {})
        return bool(context.get(chart_key, {}).get("include", True))

    def _insert_chart_limitation_after(self, target_para, note: str):
        if not note:
            return target_para
        p = self.doc.add_paragraph()
        p.text = note
        target_para._p.addnext(p._p)
        return p

    def _find_chart_anchor(self, anchor_text: str, values: Dict[str, str]):
        """Return the paragraph to insert a chart after."""
        needle = anchor_text.strip().lower()
        for para in self.doc.paragraphs:
            if needle and needle in para.text.strip().lower():
                return para

        # Fallbacks for generated prose replacing the template placeholders.
        for key in ("G_TREND_NARRATIVE", "G_TREND_REPORTS_SUMMARY", "C_SALES_METHOD"):
            value = values.get(key, "")
            if not value:
                continue
            prefix = value[:60].strip().lower()
            for para in self.doc.paragraphs:
                if prefix and prefix in para.text.strip().lower():
                    return para
        return None

    def _insert_chart_after(self, target_para, chart_path: Path, caption: str):
        """Insert caption and chart after target_para and return the chart paragraph."""
        caption_para = self.doc.add_paragraph()
        caption_para.text = caption
        target_para._p.addnext(caption_para._p)

        chart_para = self.doc.add_paragraph()
        caption_para._p.addnext(chart_para._p)
        run = chart_para.add_run()
        run.add_picture(str(chart_path), width=Inches(6))
        note = self._chart_note_from_caption(caption)
        if note:
            note_para = self.doc.add_paragraph()
            note_para.text = note
            chart_para._p.addnext(note_para._p)
            return note_para
        return chart_para

    def _insert_picture_after(self, target_para, chart_path: Path, note: str = ""):
        chart_para = self.doc.add_paragraph()
        target_para._p.addnext(chart_para._p)
        run = chart_para.add_run()
        run.add_picture(str(chart_path), width=Inches(6))
        if note:
            note_para = self.doc.add_paragraph()
            note_para.text = note
            chart_para._p.addnext(note_para._p)
            return note_para
        return chart_para

    def _chart_note_from_caption(self, caption: str) -> str:
        for chart_key, data in getattr(self, "_chart_context", {}).items():
            if data.get("caption") == caption:
                return data.get("note", "")
        return ""

    def _build_chart_context(self, psur: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        stats = psur.get("_statistics", {}) or {}

        def num(v, default=0):
            try:
                return float(v)
            except (TypeError, ValueError):
                return default

        total_units = int(num(stats.get("total_units_sold")))
        total_complaints = int(num(stats.get("total_complaints")))
        overall_pct = num(stats.get("overall_complaint_percentage"))
        if not overall_pct and total_units:
            overall_pct = round(total_complaints / total_units * 100, 4)

        region_rows = stats.get("section_c_region_rows") or []
        region_bits = []
        for row in region_rows:
            if not isinstance(row, dict) or row.get("region") == "Worldwide":
                continue
            units = int(num(row.get("units")))
            if units <= 0:
                continue
            pct = num(row.get("pct_current"))
            region_bits.append(f"{row.get('region')}: {units:,} units ({pct:.1f}%)")

        rates_by_harm = stats.get("rates_by_harm") or []
        top_harm = max(rates_by_harm, key=lambda r: r.get("complaint_count", 0), default={})
        harm_note = (
            f"Interpretation: {int(top_harm.get('complaint_count', 0))} of {total_complaints} complaints "
            f"({num(top_harm.get('percentage')):.4f}% of distributed units) were classified as "
            f"{top_harm.get('category')}; this keeps the complaint profile anchored to the same denominator "
            "used in Table 7."
            if top_harm else
            "Interpretation: complaint harm categories were not available for charting."
        )

        rates_by_imdrf = stats.get("rates_by_imdrf") or []
        top_mdp = max(rates_by_imdrf, key=lambda r: r.get("complaint_count", 0), default={})
        mdp_note = (
            f"Interpretation: the leading medical device problem was {top_mdp.get('category')} "
            f"with {int(top_mdp.get('complaint_count', 0))} complaints "
            f"({num(top_mdp.get('percentage')):.4f}% of distributed units), which is the primary driver "
            "to consider when reconciling complaint themes to the RMF and CAPA discussion."
            if top_mdp else
            "Interpretation: medical device problem categories were not available for charting."
        )

        trend = stats.get("trend_analysis") or {}
        trend_status = str(trend.get("status") or "N/A")
        mean_pct = num(trend.get("mean_pct"))
        ucl_pct = num(trend.get("ucl_3sigma_pct"))
        current_pct = num(trend.get("current_rate_pct"))
        monthly = trend.get("monthly_rates_pct") or []
        max_monthly = max(monthly) if monthly else 0
        labels = trend.get("monthly_labels") or []
        peak_label = labels[monthly.index(max_monthly)] if monthly and labels else "the reporting period"

        return {
            "sales_trend": {
                "include": bool(stats.get("units_by_month")),
                "caption": "Figure C-1. Sales trend by month for the 2023 surveillance period.",
                "note": (
                    f"Interpretation: total exposure for this PSUR is {total_units:,} distributed units"
                    + (f", split across {', '.join(region_bits)}." if region_bits else ".")
                    + " This denominator is used consistently for complaint and serious-incident rate calculations."
                ),
            },
            "complaints_region": {
                "include": bool(stats.get("complaints_by_region")),
                "caption": "Figure F-1. Complaints by region.",
                "note": "Regional complaint attribution was not available in the parsed complaint dataset, so regional complaint distribution is not charted; Table 1 still provides the regional exposure denominator.",
            },
            "harm_distribution": {
                "include": bool(stats.get("complaints_by_harm")),
                "caption": "Figure F-1. Complaint distribution by IMDRF harm category.",
                "note": harm_note,
            },
            "top_mdps": {
                "include": bool(stats.get("complaints_by_imdrf")),
                "caption": "Figure F-2. Top IMDRF medical device problem categories.",
                "note": mdp_note,
            },
            "trend_ucl": {
                "include": bool(monthly),
                "caption": "Figure G-1. Monthly complaint rate control chart.",
                "note": (
                    f"Interpretation: the monthly complaint-rate profile is {trend_status}; mean monthly rate was "
                    f"{mean_pct:.2f}%, UCL was {ucl_pct:.2f}%, and the period-end monthly rate was {current_pct:.2f}%."
                ),
            },
            "rate_occurrence": {
                "include": bool(monthly),
                "caption": "Figure G-2. Monthly complaint rate against occurrence reference bands.",
                "note": (
                    f"Interpretation: the highest monthly rate was {max_monthly:.2f}% in {peak_label}; this view shows "
                    "whether observed complaint rates remain within the expected occurrence band used for risk trending."
                ),
            },
            "harm_trend": {
                "include": bool(stats.get("harm_by_month")),
                "caption": "Figure G-3. Monthly complaint count by harm category.",
                "note": "Interpretation: the time-series view separates no-health-consequence complaints from laceration complaints so any change in harm severity can be assessed separately from complaint volume.",
            },
            "per_period": {
                "include": bool(stats.get("per_period_aggregates")),
                "caption": "Figure G-4. Complaint count and rate by reporting period.",
                "note": "Prior comparable 12-month aggregate data were not available, so the report does not include a year-over-year period chart; the current-period monthly control chart is used for trend assessment.",
            },
            "ract_matrix": {
                "include": bool(stats.get("ract_matrix") or stats.get("risk_summary")),
                "caption": "Figure G-4. RACT severity and occurrence matrix.",
                "note": "RACT severity/occurrence matrix data were not available in the parsed inputs, so this chart is not included; Table 7 still identifies where RACT thresholds were not provided.",
            },
        }

    def _replace_in_runs(self, para, values: Dict[str, str]):
        """Replace placeholders across runs, handling split markers."""
        runs = para.runs
        if not runs:
            return

        full = "".join(r.text for r in runs)
        if "{{" not in full:
            return

        new_full = PH_RE.sub(lambda m: values.get(m.group(1), m.group(0)), full)

        if new_full == full:
            return

        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ""

    def _update_headers(self, psur, values: Dict[str, str]):
        """Fill header/footer placeholders with product info."""
        if not self.doc or not self.doc.sections:
            return

        for section in self.doc.sections:
            for para in section.header.paragraphs:
                if "{{" in para.text:
                    self._replace_in_runs(para, values)
            for para in section.footer.paragraphs:
                if "{{" in para.text:
                    self._replace_in_runs(para, values)
