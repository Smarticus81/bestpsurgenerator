"""FormattingMixin — universal DOCX formatting enforcement for the PSUR renderer.

Applies consistent typography rules across the entire document:
  - Titles (Heading 1):   Arial 12pt, Bold, Underlined
  - Subtitles (Heading 2/3): Arial 10pt, Bold (no underline)
  - Body / everything else: Arial 10pt, non-bold, non-underlined
  - Field labels (bold) vs values (non-bold) in colon-separated paragraphs

Also provides insertion of the synthesized Executive Summary narrative
and Benefit-Risk Conclusion Statement into Section A.
"""
import logging
import re
from typing import Dict

from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Font constants
FONT_NAME = "Arial"
TITLE_SIZE = Pt(12)
SUBTITLE_SIZE = Pt(10)
BODY_SIZE = Pt(10)

# Heading-style patterns (case-insensitive)
_TITLE_STYLES = {"Heading 1", "Heading1", "Title", "heading 1"}
_SUBTITLE_STYLES = {"Heading 2", "Heading2", "Heading 3", "Heading3",
                     "Subtitle", "heading 2", "heading 3"}

# Section heading regex: "Section A:", "A.", "Section B:", etc.
_SECTION_HEADING_RE = re.compile(
    r"^(Section\s+)?[A-M][\.\:]\s", re.IGNORECASE
)

# Sub-section label patterns (typical FormQAR-054 sub-titles)
_SUBTITLE_TEXT_PATTERNS = [
    re.compile(r"^[A-M]\.\d+", re.IGNORECASE),           # "A.1 ..."
    re.compile(r"^(Previous PSUR|Notified Body|Data Collection|"
               r"Benefit.Risk|Device Information|Device Classification|"
               r"Device Description|Intended Purpose|Clinical Data|"
               r"Serious Incident|Customer Feedback|Complaint|"
               r"Trend|Corrective|Preventive|FSCA|CAPA|PMCF|"
               r"Review of External|Findings|Overall|"
               r"Volume of Sales|Population Exposure|"
               r"Manufacturer Information|Regulatory Information|"
               r"Document Information)", re.IGNORECASE),
]

# Checkbox font (kept distinct)
CB_FONT = "Segoe UI Symbol"
_CHECKBOX_CHARS = {"\u2610", "\u2611", "\u2612"}


class FormattingMixin:
    """Provides universal DOCX formatting enforcement and Section A content insertion."""

    # ==================================================================
    # Section A content insertion
    # ==================================================================
    def _insert_section_a_content(self, values: Dict[str, str]):
        """Insert synthesized Executive Summary narrative and B-R Conclusion
        Statement into Section A of the document if not already placed via
        template placeholders.
        """
        exec_narrative = values.get("A_EXECUTIVE_SUMMARY_NARRATIVE", "")
        brc_statement = values.get("A_BRC_CONCLUSION_STATEMENT", "")

        if not exec_narrative and not brc_statement:
            return

        # Detect whether they were already placed by placeholder substitution
        exec_placed = False
        brc_placed = False
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if exec_narrative and exec_narrative[:60] in text:
                exec_placed = True
            if brc_statement and brc_statement[:60] in text:
                brc_placed = True

        # Find Section A area
        section_a_start = None
        section_b_start = None
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip().upper()
            if ("SECTION A" in text or "EXECUTIVE SUMMARY" in text) and section_a_start is None:
                section_a_start = i
            elif section_a_start is not None and (
                "SECTION B" in text or "SCOPE AND DEVICE" in text
            ):
                section_b_start = i
                break

        if section_a_start is None:
            logger.warning("Could not locate Section A in document for content insertion")
            return

        # Insert Executive Summary Narrative right after the Section A heading
        if exec_narrative and not exec_placed:
            insert_idx = section_a_start
            target_para = self.doc.paragraphs[insert_idx]
            new_elem = OxmlElement("w:p")
            target_para._element.addnext(new_elem)

            # Re-fetch to get wrapped paragraph
            for p in self.doc.paragraphs:
                if p._element is new_elem:
                    run = p.add_run(exec_narrative)
                    run.font.name = FONT_NAME
                    run.font.size = BODY_SIZE
                    run.font.bold = False
                    run.font.italic = False
                    run.font.underline = False
                    logger.info("Inserted Executive Summary narrative after Section A heading")
                    break

        # Insert B-R Conclusion Statement before Section B
        # (or at the end of Section A area)
        if brc_statement and not brc_placed:
            # Find the best insertion point: after the last B-R checkbox area
            insert_idx = None

            # Re-fetch paragraphs since we may have inserted one above
            paras = self.doc.paragraphs
            for i, para in enumerate(paras):
                text = para.text.strip()
                # Look for the B-R assessment conclusion area markers
                if ("ADVERSELY_IMPACTED" in text.upper()
                        or "NOT ADVERSELY IMPACTED" in text.upper()
                        or "\u2611" in text or "\u2610" in text):
                    # Check if this is in Section A area
                    if section_a_start is not None and i >= section_a_start:
                        if section_b_start is None or i < section_b_start + 5:
                            insert_idx = i

                # Also match the A_BRC_SUMMARY placeholder area
                if "A_BRC_SUMMARY" in text or "high_level_summary" in text.lower():
                    if section_a_start is not None and i >= section_a_start:
                        insert_idx = i

            if insert_idx is None:
                # Fallback: insert before Section B start or at end of Section A
                if section_b_start is not None:
                    insert_idx = section_b_start - 1
                else:
                    insert_idx = section_a_start + 3  # after heading + some content

            target_para = paras[min(insert_idx, len(paras) - 1)]

            # Add subtitle "Benefit-Risk Analysis Conclusion"
            subtitle_elem = OxmlElement("w:p")
            target_para._element.addnext(subtitle_elem)
            for p in self.doc.paragraphs:
                if p._element is subtitle_elem:
                    run = p.add_run("Benefit-Risk Analysis Conclusion")
                    run.font.name = FONT_NAME
                    run.font.size = SUBTITLE_SIZE
                    run.font.bold = True
                    run.font.italic = False
                    run.font.underline = False
                    break

            # Add the conclusion statement paragraph after the subtitle
            stmt_elem = OxmlElement("w:p")
            subtitle_elem.addnext(stmt_elem)
            for p in self.doc.paragraphs:
                if p._element is stmt_elem:
                    run = p.add_run(brc_statement)
                    run.font.name = FONT_NAME
                    run.font.size = BODY_SIZE
                    run.font.bold = False
                    run.font.italic = False
                    run.font.underline = False
                    logger.info("Inserted B-R Conclusion Statement in Section A")
                    break

    # ==================================================================
    # Universal formatting enforcement
    # ==================================================================
    def _apply_universal_formatting(self):
        """Walk every paragraph and run in the document, enforcing universal
        typography rules:
          - Titles:    Arial 12pt Bold Underlined
          - Subtitles: Arial 10pt Bold (no underline)
          - Body text: Arial 10pt non-bold non-underlined
          - Field labels (before ':') bold, values (after ':') non-bold
        """
        # Clear italic from document default style if present
        try:
            default_style = self.doc.styles["Normal"]
        except (KeyError, TypeError):
            default_style = None
        if default_style and default_style.font:
            default_style.font.italic = False
            default_style.font.name = FONT_NAME

        for para in self.doc.paragraphs:
            role = self._classify_paragraph(para)
            if role == "title":
                self._format_title(para)
            elif role == "subtitle":
                self._format_subtitle(para)
            else:
                self._format_body(para)

        # Also format table cells (body style)
        for table in self.doc.tables:
            self._format_table_cells(table)

        # Also format headers and footers
        for section in self.doc.sections:
            for hf in (section.header, section.footer,
                       section.first_page_header, section.first_page_footer,
                       section.even_page_header, section.even_page_footer):
                if hf is None:
                    continue
                for para in hf.paragraphs:
                    for run in para.runs:
                        if self._is_checkbox_run(run):
                            continue
                        run.font.name = FONT_NAME
                        run.font.italic = False
                        self._set_east_asia_font(run, FONT_NAME)

    # ------------------------------------------------------------------
    # Paragraph classification
    # ------------------------------------------------------------------
    def _classify_paragraph(self, para) -> str:
        """Return 'title', 'subtitle', or 'body' for a paragraph."""
        style_name = (para.style.name or "").strip() if para.style else ""

        # Style-based classification
        if style_name in _TITLE_STYLES:
            return "title"
        if style_name in _SUBTITLE_STYLES:
            return "subtitle"

        # Content-based classification
        text = para.text.strip()
        if not text:
            return "body"

        # Section headings like "Section A: Executive Summary"
        if _SECTION_HEADING_RE.match(text):
            return "title"

        # Label/value fields such as "Device Description: ..." and
        # "Intended Purpose/Use: ..." are body paragraphs. Only the label may
        # be bold; the content after the colon must remain regular weight.
        colon_idx = text.find(":")
        if 2 < colon_idx < 80 and colon_idx < len(text) - 1:
            label = text[:colon_idx].strip().lower()
            if label in {"device description", "intended purpose", "intended purpose/use", "intended use"}:
                return "body"

        # Sub-section patterns
        for pat in _SUBTITLE_TEXT_PATTERNS:
            if pat.match(text):
                return "subtitle"

        return "body"

    # ------------------------------------------------------------------
    # Formatting helpers
    # ------------------------------------------------------------------
    def _format_title(self, para):
        """Apply Title formatting: Arial 12pt Bold Underlined."""
        for run in para.runs:
            if self._is_checkbox_run(run):
                continue
            run.font.name = FONT_NAME
            run.font.size = TITLE_SIZE
            run.font.bold = True
            run.font.italic = False
            run.font.underline = True
            self._set_east_asia_font(run, FONT_NAME)

    def _format_subtitle(self, para):
        """Apply Subtitle formatting: Arial 10pt Bold (no underline)."""
        for run in para.runs:
            if self._is_checkbox_run(run):
                continue
            run.font.name = FONT_NAME
            run.font.size = SUBTITLE_SIZE
            run.font.bold = True
            run.font.italic = False
            run.font.underline = False
            self._set_east_asia_font(run, FONT_NAME)

    def _format_body(self, para):
        """Apply Body formatting: Arial 10pt, non-bold, non-underlined.

        Special handling: if the paragraph contains a colon-separated
        field label and value, the label portion is set to bold.
        """
        full_text = para.text
        colon_idx = full_text.find(":")

        # Determine if this is a "Field: Value" pattern
        # Heuristic: colon within the first 80 chars, at least 3 chars before it,
        # and text after the colon
        is_field_value = (
            colon_idx > 2
            and colon_idx < 80
            and colon_idx < len(full_text) - 1
            and not full_text[:colon_idx].strip().startswith("http")
            and "\n" not in full_text[:colon_idx]
        )

        if is_field_value and len(para.runs) >= 1:
            self._format_field_value_paragraph(para, colon_idx)
        else:
            for run in para.runs:
                if self._is_checkbox_run(run):
                    continue
                run.font.name = FONT_NAME
                run.font.size = BODY_SIZE
                run.font.bold = False
                run.font.italic = False
                run.font.underline = False
                self._set_east_asia_font(run, FONT_NAME)

    def _format_field_value_paragraph(self, para, colon_idx: int):
        """Split a paragraph into bold field label and non-bold value at the colon."""
        full_text = para.text
        label_text = full_text[:colon_idx + 1]  # Include the colon
        value_text = full_text[colon_idx + 1:]

        # Clear all existing runs
        runs = para.runs
        if not runs:
            return

        # Reconstruct with two runs: bold label, non-bold value
        # Use the first run for the label
        runs[0].text = label_text
        runs[0].font.name = FONT_NAME
        runs[0].font.size = BODY_SIZE
        runs[0].font.bold = True
        runs[0].font.italic = False
        runs[0].font.underline = False
        self._set_east_asia_font(runs[0], FONT_NAME)

        if len(runs) > 1:
            # Use second run for value
            runs[1].text = value_text
            runs[1].font.name = FONT_NAME
            runs[1].font.size = BODY_SIZE
            runs[1].font.bold = False
            runs[1].font.italic = False
            runs[1].font.underline = False
            self._set_east_asia_font(runs[1], FONT_NAME)
            # Clear remaining runs
            for r in runs[2:]:
                r.text = ""
        else:
            # Only one run: need to add a second for the value
            new_run = para.add_run(value_text)
            new_run.font.name = FONT_NAME
            new_run.font.size = BODY_SIZE
            new_run.font.bold = False
            new_run.font.italic = False
            new_run.font.underline = False
            self._set_east_asia_font(new_run, FONT_NAME)
            runs[0].text = label_text  # Trim label from first run

    def _format_table_cells(self, table):
        """Apply body formatting to all table cells (Arial 10, non-underlined).

        Header rows (row 0) get bold; data rows get non-bold.
        """
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if self._is_checkbox_run(run):
                            continue
                        run.font.name = FONT_NAME
                        run.font.size = BODY_SIZE
                        run.font.italic = False
                        run.font.underline = False
                        self._set_east_asia_font(run, FONT_NAME)
                        # Header row gets bold, data rows don't
                        if ri == 0:
                            run.font.bold = True
                        else:
                            run.font.bold = False

    # ------------------------------------------------------------------
    # Utility
    # ------------------------------------------------------------------
    @staticmethod
    def _is_checkbox_run(run) -> bool:
        """Return True if the run contains only checkbox characters."""
        text = run.text.strip()
        return bool(text) and all(ch in _CHECKBOX_CHARS for ch in text)

    @staticmethod
    def _set_east_asia_font(run, font_name: str):
        """Set the East Asia font via XML to ensure full coverage."""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)
