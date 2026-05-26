"""
PSURTemplateRenderer — thin facade composing rendering mixins.

Clones FormQAR-054_template.docx and fills it from PSUR JSON data.
"""
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document

from config import DOCX_TEMPLATE_PATH
from rendering._value_map import ValueMapMixin
from rendering._paragraphs import ParagraphMixin
from rendering._tables import TableMixin
from rendering._formatting import FormattingMixin

logger = logging.getLogger(__name__)

# Regex to match template instruction debris: [bracketed text]
# Excludes checkbox characters (☐☑☒) which use similar formatting
_TEMPLATE_DEBRIS_RE = re.compile(r'\[(?![☐☑☒])[^\[\]]{3,}\]')

# Specific template instruction phrases to strip
_TEMPLATE_PHRASES = [
    "(Remove if not applicable)",
    "(remove if not applicable)",
    "[Add or delete rows as needed]",
    "[add or delete rows as needed]",
    "[Space for line chart showing:]",
    "[Note: Multiply number of sales units",
    "[Any other countries which have more than 5% of global sales. Add rows as needed.]",
    "[Any other countries which have more than 5%",
]

_CHART_INSTRUCTION_PARAGRAPHS = {
    "space for line chart showing:",
    "x-axis: time periods",
    "y-axis: units sold",
    "trend line showing overall growth/decline pattern",
}


class PSURTemplateRenderer(ValueMapMixin, ParagraphMixin, TableMixin, FormattingMixin):
    """Clone the FormQAR-054 DOCX template and fill it from PSUR JSON."""

    def __init__(self, template_path: Optional[Path] = None):
        self.template_path = Path(template_path) if template_path else DOCX_TEMPLATE_PATH
        self.doc: Optional[Document] = None
        self.chart_paths: Dict[str, Path] = {}

    def render(self, psur: Dict[str, Any], output_path: Path,
               chart_paths: Optional[Dict[str, Path]] = None,
               tables_docx_path: Optional[Path] = None) -> None:
        """Clone the template, fill all placeholders, save to *output_path*.

        If ``tables_docx_path`` is provided and exists, the data tables
        (sales / IMDRF / complaint / FSCA / CAPA / external DB / PMCF /
        health impact / customer feedback) in the rendered document are
        replaced with the deterministic standalone equivalents.
        """
        output_path = Path(output_path)
        if not self.template_path.exists():
            raise FileNotFoundError(
                f"Template not found: {self.template_path}\n"
                "Ensure FormQAR-054_template.docx is in the constraints/ folder."
            )

        self.doc = Document(str(self.template_path))
        props = self.doc.core_properties
        props.author = "Mastropietro Company Regulatory Affairs"
        props.last_modified_by = "Mastropietro Company Regulatory Affairs"
        props.title = "Periodic Safety Update Report"
        self.chart_paths = chart_paths or {}
        self._placed_chart_keys = set()
        self._chart_context = self._build_chart_context(psur)

        values = self._build_value_map(psur)

        # 0) Determine cadence and delete unused table variants
        self._delete_unused_table_variants(psur)

        # 1) Fill placeholders in body paragraphs
        for para in self.doc.paragraphs:
            self._fill_paragraph(para, values)

        # 1b) Insert charts that have no template placeholder
        self._insert_unplaced_charts(values)

        # 2) Fill data into all tables + placeholder substitution
        self._fill_all_tables(psur, values)

        # 2b) Splice deterministic standalone tables on top of filled ones
        if tables_docx_path is not None:
            self._replace_tables_from_docx(tables_docx_path)

        # 3) Update headers/footers
        self._update_headers(psur, values)

        # 4) Insert Section A synthesized content (exec summary + B-R conclusion)
        self._insert_section_a_content(values)

        # 5) Apply universal formatting (fonts, sizes, bold/underline rules)
        self._apply_universal_formatting()

        # 6) Strip all template debris (bracketed instructions, footnotes, etc.)
        self._strip_template_debris()

        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))

    # ==================================================================
    # Template debris stripping
    # ==================================================================
    def _strip_template_debris(self):
        """Remove all template instruction debris from the final document.

        This runs AFTER all content is filled, ensuring no [bracketed instructions],
        (Remove if not applicable) annotations, or footnote debris survive.
        """
        debris_count = 0

        # Strip from body paragraphs
        for para in list(self.doc.paragraphs):
            if self._remove_chart_instruction_paragraph(para):
                debris_count += 1
                continue
            debris_count += self._strip_debris_from_paragraph(para)

        # Strip from table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._remove_chart_instruction_paragraph(para):
                            debris_count += 1
                            continue
                        debris_count += self._strip_debris_from_paragraph(para)

        # Strip from headers/footers
        for section in self.doc.sections:
            for hf in (section.header, section.footer,
                       section.first_page_header, section.first_page_footer):
                if hf is None:
                    continue
                for para in hf.paragraphs:
                    if self._remove_chart_instruction_paragraph(para):
                        debris_count += 1
                        continue
                    self._strip_debris_from_paragraph(para)

        if debris_count > 0:
            logger.info(f"Stripped {debris_count} template debris instances from document")

    def _remove_chart_instruction_paragraph(self, para) -> bool:
        """Remove the FormQAR chart placeholder instruction block."""
        text = " ".join((para.text or "").replace("[", "").replace("]", "").split())
        text_lower = text.lower().lstrip("-• ").strip()
        if text_lower in _CHART_INSTRUCTION_PARAGRAPHS:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)
                return True
        return False

    def _strip_debris_from_paragraph(self, para) -> int:
        """Strip template debris from a single paragraph. Returns count of removals."""
        full_text = para.text
        if not full_text:
            return 0

        cleaned = full_text
        count = 0

        # Strip specific known phrases first
        for phrase in _TEMPLATE_PHRASES:
            if phrase in cleaned:
                cleaned = cleaned.replace(phrase, "")
                count += 1

        # Strip remaining [bracketed instructions] (but not checkbox chars)
        matches = _TEMPLATE_DEBRIS_RE.findall(cleaned)
        for match in matches:
            # Don't strip things that look like real content (e.g. [PMCF-001])
            match_lower = match.lower()
            if any(kw in match_lower for kw in (
                "use this table", "add or delete", "remove if",
                "note:", "any other countries", "multiply number",
                "select one", "delete the", "insert",
                "add rows", "complete this", "fill in",
            )):
                cleaned = cleaned.replace(match, "")
                count += 1

        if count > 0 and cleaned != full_text:
            # Apply the cleaned text back to the paragraph runs
            cleaned = cleaned.strip()
            runs = para.runs
            if runs:
                runs[0].text = cleaned
                for r in runs[1:]:
                    r.text = ""

        # Remove paragraphs that are now entirely empty after stripping
        # (only if they were ONLY template debris)
        if count > 0 and not cleaned.strip():
            # Remove the paragraph element from the document body
            para._element.getparent().remove(para._element)

        return count

    # ==================================================================
    # Delete unused table variants
    # ==================================================================
    def _delete_unused_table_variants(self, psur: Dict[str, Any]):
        """Delete the unused annual/biennial table variant from the template.

        FormQAR-054 contains BOTH annual and biennial variants for Table 1
        and Table 7. We determine cadence and remove the wrong one.
        """
        # Determine cadence from PSUR data
        sections = psur.get("sections", {})
        sec_b = sections.get("B_scope_and_device_description", {})
        classification = sec_b.get("device_classification", {})
        eu_class = (classification.get("eu_mdr_classification", "") or "").upper()

        # Class IIb, III, implantable = annual; Class IIa = biennial
        # Default to annual if classification is unclear
        is_annual = eu_class not in ("CLASS_IIA",)

        # Also check the doc_info cadence field
        cover = psur.get("psur_cover_page", {})
        doc_info = cover.get("document_information", {})
        cadence = (doc_info.get("psur_cadence", "") or "").lower()
        if "biennial" in cadence or "24" in cadence:
            is_annual = False
        elif "annual" in cadence or "12" in cadence:
            is_annual = True

        # Keywords that identify variant tables
        annual_keywords = ["annually", "annual number", "12-month", "12 month"]
        biennial_keywords = ["every two years", "biennial", "24-month", "24 month"]
        remove_keywords = biennial_keywords if is_annual else annual_keywords

        # Remove matching paragraphs (variant headers/instructions)
        paras_to_remove = []
        for para in self.doc.paragraphs:
            text = para.text.lower()
            if any(kw in text for kw in remove_keywords):
                # Check it's a table instruction, not real content
                if any(marker in text for marker in [
                    "use this table", "select this", "[",
                    "delete the other", "table if",
                ]):
                    paras_to_remove.append(para)

        for para in paras_to_remove:
            para._element.getparent().remove(para._element)

        logger.info(
            f"Cadence: {'annual' if is_annual else 'biennial'}. "
            f"Removed {len(paras_to_remove)} unused variant instruction paragraphs."
        )


# ═════════════════════════════════════════════════════════════════════
# CLI ENTRY POINT
# ═════════════════════════════════════════════════════════════════════

def main():
    """CLI: render a PSUR JSON into a filled DOCX."""
    import argparse
    parser = argparse.ArgumentParser(description="Render PSUR JSON into FormQAR-054 template")
    parser.add_argument("json_path", help="Path to PSUR JSON file")
    parser.add_argument("-o", "--output", default="PSUR_RENDERED.docx", help="Output DOCX path")
    parser.add_argument("-t", "--template", default=None, help="Template DOCX path")
    parser.add_argument("--sales-chart", default=None, help="Sales trend chart image path")
    parser.add_argument("--trend-chart", default=None, help="UCL trend chart image path")
    args = parser.parse_args()

    with open(args.json_path) as f:
        psur = json.load(f)

    charts = {}
    if args.sales_chart:
        charts["sales_trend"] = Path(args.sales_chart)
    if args.trend_chart:
        charts["trend_ucl"] = Path(args.trend_chart)

    renderer = PSURTemplateRenderer(template_path=args.template)
    renderer.render(psur, Path(args.output), chart_paths=charts)
    print(f"Rendered: {args.output}")


if __name__ == "__main__":
    main()
