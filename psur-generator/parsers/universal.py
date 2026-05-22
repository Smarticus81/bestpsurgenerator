"""Universal multi-format file parser.

Handles: .csv, .xlsx, .xls, .docx, .pdf, .json, .md, .txt, and images (.png, .jpg, .jpeg, .tiff, .bmp)
Images are processed via Claude Vision API.
Includes encoding auto-detection, markdown table parsing, and content-based format fallback.
"""
import json
import base64
import logging
import mimetypes
import re
from pathlib import Path
from typing import Any, Dict, Optional

from config import MODEL
from llm_client import get_llm_client

logger = logging.getLogger(__name__)

# Extensions by category
TABULAR_EXTENSIONS = {".csv", ".xlsx", ".xls"}
DOCUMENT_EXTENSIONS = {".docx", ".doc"}
PDF_EXTENSIONS = {".pdf"}
TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
JSON_EXTENSIONS = {".json"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif", ".webp"}


def detect_format(filepath: Path) -> str:
    """Detect file format category. Falls back to content-based detection."""
    ext = filepath.suffix.lower()
    if ext in TABULAR_EXTENSIONS:
        return "tabular"
    elif ext in DOCUMENT_EXTENSIONS:
        return "document"
    elif ext in PDF_EXTENSIONS:
        return "pdf"
    elif ext in TEXT_EXTENSIONS:
        return "text"
    elif ext in JSON_EXTENSIONS:
        return "json"
    elif ext in IMAGE_EXTENSIONS:
        return "image"
    else:
        # Content-based fallback detection
        detected = _detect_format_by_content(filepath)
        if detected:
            logger.info(f"Detected {filepath.name} as '{detected}' via content inspection")
            return detected
        raise ValueError(f"Unsupported file format: {ext} for {filepath.name}")


def _detect_format_by_content(filepath: Path) -> Optional[str]:
    """Detect file format by inspecting file content (magic bytes / structure)."""
    try:
        header = filepath.read_bytes()[:16]
    except Exception:
        return None

    # PDF magic: %PDF-
    if header[:5] == b"%PDF-":
        return "pdf"
    # DOCX/XLSX are ZIP files: PK\x03\x04
    if header[:4] == b"PK\x03\x04":
        # Peek inside to distinguish DOCX from XLSX
        import zipfile
        try:
            with zipfile.ZipFile(filepath) as zf:
                names = zf.namelist()
                if any("word/" in n for n in names):
                    return "document"
                if any("xl/" in n for n in names):
                    return "tabular"
        except Exception:
            pass
        return "document"  # Default ZIP-based to document
    # XLS magic: \xD0\xCF\x11\xE0 (OLE compound)
    if header[:4] == b"\xd0\xcf\x11\xe0":
        return "tabular"
    # Try reading as text to check for CSV/JSON
    try:
        text = filepath.read_text(encoding="utf-8", errors="replace")[:2000]
        text_stripped = text.strip()
        if text_stripped.startswith("{") or text_stripped.startswith("["):
            return "json"
        # If it looks like CSV (has commas and newlines)
        lines = text_stripped.split("\n")
        if len(lines) > 1 and all("," in line for line in lines[:5]):
            return "tabular"
        # Otherwise treat as text
        if text_stripped:
            return "text"
    except Exception:
        pass

    return None


def parse_file(filepath: Path, purpose: str = "", date_filter: Optional[Dict] = None) -> Dict[str, Any]:
    """
    Parse any supported file into a standardized dict.

    Args:
        filepath: Path to the file
        purpose: Hint about what data this file contains (e.g. "sales", "complaints")
        date_filter: Optional {"start": "YYYY-MM-DD", "end": "YYYY-MM-DD"}

    Returns:
        Dict with parsed content and metadata
    """
    filepath = Path(filepath)
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    fmt = detect_format(filepath)

    result = {
        "source_file": filepath.name,
        "file_format": fmt,
        "file_extension": filepath.suffix.lower(),
        "purpose": purpose,
    }

    if fmt == "tabular":
        result.update(_parse_tabular(filepath, date_filter))
    elif fmt == "document":
        result.update(_parse_docx(filepath))
    elif fmt == "pdf":
        result.update(_parse_pdf(filepath))
    elif fmt == "text":
        result.update(_parse_text(filepath))
    elif fmt == "json":
        result.update(_parse_json(filepath))
    elif fmt == "image":
        result.update(_parse_image(filepath, purpose))

    return result


def _parse_tabular(filepath: Path, date_filter: Optional[Dict] = None) -> Dict[str, Any]:
    """Parse CSV or Excel into rows + metadata with encoding auto-detection."""
    import pandas as pd

    ext = filepath.suffix.lower()
    if ext == ".csv":
        df = _read_csv_with_encoding(filepath)
    else:
        df = pd.read_excel(filepath)

    # Normalize column names
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]

    return {
        "type": "tabular",
        "columns": list(df.columns),
        "row_count": len(df),
        "dataframe": df,  # kept for downstream parsers
        "preview": df.head(5).to_dict(orient="records"),
    }


def _read_csv_with_encoding(filepath: Path) -> "pd.DataFrame":
    """Read CSV with automatic encoding detection. Tries multiple encodings."""
    import pandas as pd

    encodings = ["utf-8", "utf-8-sig", "latin-1", "iso-8859-1", "cp1252", "utf-16"]

    for encoding in encodings:
        try:
            return pd.read_csv(filepath, encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            # If it's not an encoding error, re-raise
            if "codec" in str(e).lower() or "decode" in str(e).lower():
                continue
            raise

    # Last resort: read with error replacement
    logger.warning(f"Could not detect encoding for {filepath.name}, using utf-8 with error replacement")
    return pd.read_csv(filepath, encoding="utf-8", errors="replace")


def _parse_docx(filepath: Path) -> Dict[str, Any]:
    """Parse Word document to text."""
    from docx import Document
    doc = Document(str(filepath))

    paragraphs = []
    tables = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                "text": text,
                "style": para.style.name if para.style else "",
            })

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    full_text = "\n".join(p["text"] for p in paragraphs)

    return {
        "type": "document",
        "full_text": full_text,
        "paragraphs": paragraphs,
        "tables": tables,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
    }


def _parse_pdf(filepath: Path) -> Dict[str, Any]:
    """Parse PDF to text and tables."""
    import pdfplumber

    full_text = ""
    page_texts = []
    tables = []

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            page_texts.append(text)
            full_text += text + "\n"

            page_tables = page.extract_tables()
            if page_tables:
                tables.extend(page_tables)

    return {
        "type": "pdf",
        "full_text": full_text,
        "page_count": len(page_texts),
        "page_texts": page_texts,
        "tables": tables,
        "table_count": len(tables),
    }


def _parse_text(filepath: Path) -> Dict[str, Any]:
    """Parse plain text or markdown. Extracts markdown tables if present."""
    text = filepath.read_text(encoding="utf-8", errors="replace")

    result = {
        "type": "text",
        "full_text": text,
        "line_count": text.count("\n") + 1,
        "char_count": len(text),
    }

    # Extract markdown tables if present
    md_tables = _extract_markdown_tables(text)
    if md_tables:
        result["tables"] = md_tables
        result["table_count"] = len(md_tables)

    return result


def _extract_markdown_tables(text: str) -> list:
    """Extract tables from markdown-formatted text."""
    tables = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # Detect markdown table: line with pipes and next line is separator
        if "|" in line and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.match(r"^[\|\s\-:]+$", next_line):
                # Found a markdown table — parse it
                table_rows = []
                # Header row
                header = [c.strip() for c in line.split("|") if c.strip()]
                table_rows.append(header)
                i += 2  # Skip separator line
                # Data rows
                while i < len(lines) and "|" in lines[i]:
                    row = [c.strip() for c in lines[i].split("|") if c.strip()]
                    if row:
                        table_rows.append(row)
                    i += 1
                if len(table_rows) > 1:
                    tables.append(table_rows)
                continue
        i += 1
    return tables


def _parse_json(filepath: Path) -> Dict[str, Any]:
    """Parse JSON file."""
    with open(filepath) as f:
        data = json.load(f)
    return {
        "type": "json",
        "data": data,
    }


def _parse_image(filepath: Path, purpose: str = "") -> Dict[str, Any]:
    """Parse image using Vision API to extract text/data."""

    client = get_llm_client()

    # Read and encode image
    image_data = filepath.read_bytes()
    base64_image = base64.b64encode(image_data).decode("utf-8")

    # Detect media type
    media_type = mimetypes.guess_type(str(filepath))[0] or "image/png"

    prompt = f"""Extract all text, data, tables, and structured information from this image.
Purpose context: {purpose or 'medical device regulatory document'}

Return a structured extraction with:
1. All visible text (preserve structure)
2. Any tables (as rows/columns)
3. Any numerical data
4. Any relevant metadata

Format as structured text."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64_image,
                    }
                },
                {"type": "text", "text": prompt}
            ]
        }]
    )

    extracted_text = response.content[0].text

    return {
        "type": "image",
        "extracted_text": extracted_text,
        "media_type": media_type,
        "file_size": len(image_data),
    }


def parse_any_to_text(filepath: Path, purpose: str = "") -> str:
    """Parse any file and return just the text content. Convenience wrapper."""
    result = parse_file(filepath, purpose=purpose)
    fmt = result.get("type", "")

    if fmt == "tabular":
        df = result.get("dataframe")
        if df is not None:
            return df.to_string()
        return str(result.get("preview", ""))
    elif fmt in ("document", "pdf", "text"):
        return result.get("full_text", "")
    elif fmt == "json":
        return json.dumps(result.get("data", {}), indent=2, default=str)
    elif fmt == "image":
        return result.get("extracted_text", "")
    return ""
