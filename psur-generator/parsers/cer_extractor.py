"""Comprehensive CER (Clinical Evaluation Report) data extractor.

Extracts structured regulatory, device, and clinical data from CER documents
(PDF or DOCX) using text extraction and LLM-based intelligent parsing.

Extracted categories:
- Regulatory: notified body, CE mark, classification, certificates, MDR/MDD
- Identifiers: UDI-DI, UDI-PI, GMDN, EMDN, model/catalog numbers
- Device Description: name, variants, accessories, materials, mechanism of action
- Indications: intended purpose, indications, contraindications, warnings
- Population: target patients, age range, anatomical site, clinical conditions
- IFU Information: instructions for use key content, precautions
- Safety/Efficacy: clinical evidence summary, PMCF, risk/benefit, residual risks
- Tables: all embedded tables with their section context
"""
import json
import logging
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Dict, List

from llm_client import get_llm_client
from config import MODEL

logger = logging.getLogger(__name__)


@dataclass
class ManufacturerInfo:
    """Manufacturer / legal manufacturer information extracted from CER."""
    company_name: str = ""
    address_lines: List[str] = field(default_factory=list)
    manufacturer_srn: str = ""  # e.g. "US-MF-000002607"
    country: str = ""
    authorized_representative_name: str = ""
    authorized_representative_address: List[str] = field(default_factory=list)
    authorized_representative_srn: str = ""  # e.g. "NL-AR-0000000059"


@dataclass
class RegulatoryInfo:
    """Regulatory and certification data extracted from CER."""
    notified_body_name: str = ""
    notified_body_number: str = ""
    ce_mark_status: str = ""
    classification_eu: str = ""  # I, IIa, IIb, III
    classification_us_fda: str = ""  # I, II, III
    classification_rule: str = ""  # e.g. "Rule 8", "Rule 6"
    conformity_route: str = ""  # e.g., Annex IX, Annex XI
    certificates: List[Dict[str, str]] = field(default_factory=list)
    mdr_mdd_reference: str = ""  # MDR 2017/745 or MDD 93/42/EEC
    regulatory_history: str = ""
    fda_clearance: str = ""  # e.g., 510(k) number
    eu_technical_documentation_number: str = ""  # e.g. "TD103"
    first_ce_marking_date: str = ""
    first_declaration_of_conformity_date: str = ""
    risk_management_file_number: str = ""  # e.g. "RMF-103"


@dataclass
class DeviceIdentifiers:
    """Unique device identifiers extracted from CER."""
    udi_di: str = ""
    udi_pi: str = ""
    gmdn_codes: List[Dict[str, str]] = field(default_factory=list)  # [{code, term}]
    emdn_codes: List[Dict[str, str]] = field(default_factory=list)  # [{code, term}]
    model_numbers: List[str] = field(default_factory=list)
    catalog_numbers: List[str] = field(default_factory=list)
    manufacturer_identifiers: List[str] = field(default_factory=list)


@dataclass
class DeviceDescription:
    """Device description data extracted from CER."""
    device_name: str = ""
    device_variants: List[str] = field(default_factory=list)
    accessories: List[str] = field(default_factory=list)
    components: List[str] = field(default_factory=list)
    materials: List[str] = field(default_factory=list)
    technology: str = ""  # e.g., mechanism of action
    intended_lifespan: str = ""
    sterilization: str = ""
    packaging: str = ""
    novel_features: str = ""


@dataclass
class IndicationsInfo:
    """Indications and intended purpose data from CER."""
    intended_purpose: str = ""
    medical_indications: List[str] = field(default_factory=list)
    contraindications: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    precautions: List[str] = field(default_factory=list)


@dataclass
class PopulationInfo:
    """Target patient population data from CER."""
    target_population: str = ""
    age_range: str = ""
    inclusion_criteria: List[str] = field(default_factory=list)
    exclusion_criteria: List[str] = field(default_factory=list)
    anatomical_site: str = ""
    clinical_conditions: List[str] = field(default_factory=list)
    estimated_patient_exposure: str = ""


@dataclass
class IFUInfo:
    """Instructions For Use information from CER."""
    ifu_summary: str = ""
    use_instructions: List[str] = field(default_factory=list)
    cleaning_reprocessing: str = ""
    storage_handling: str = ""
    training_requirements: str = ""


@dataclass
class SafetyEfficacy:
    """Safety and efficacy data from CER."""
    clinical_evidence_summary: str = ""
    clinical_investigations: List[str] = field(default_factory=list)
    literature_review_summary: str = ""
    equivalence_assessment: str = ""
    pmcf_requirements: str = ""
    pmcf_planned_activities: List[str] = field(default_factory=list)
    risk_benefit_conclusion: str = ""
    residual_risks: List[str] = field(default_factory=list)
    state_of_the_art: str = ""
    overall_conclusions: str = ""


@dataclass
class CERData:
    """Complete structured data extracted from a CER document."""
    manufacturer: ManufacturerInfo = field(default_factory=ManufacturerInfo)
    regulatory: RegulatoryInfo = field(default_factory=RegulatoryInfo)
    identifiers: DeviceIdentifiers = field(default_factory=DeviceIdentifiers)
    device_description: DeviceDescription = field(default_factory=DeviceDescription)
    indications: IndicationsInfo = field(default_factory=IndicationsInfo)
    population: PopulationInfo = field(default_factory=PopulationInfo)
    ifu: IFUInfo = field(default_factory=IFUInfo)
    safety_efficacy: SafetyEfficacy = field(default_factory=SafetyEfficacy)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    source_file: str = ""
    total_pages: int = 0
    extraction_method: str = ""  # "pdf" or "docx"

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    def to_flat_dict(self) -> Dict[str, Any]:
        """Flatten key fields for backward compatibility with existing CER parser."""
        flat = {
            # Backward-compatible fields
            "device_description": self.device_description.device_name
            + ("\n\n" + self.device_description.technology if self.device_description.technology else ""),
            "intended_use": self.indications.intended_purpose,
            "indications": "\n".join(self.indications.medical_indications) if self.indications.medical_indications else "",
            "contraindications": "\n".join(self.indications.contraindications) if self.indications.contraindications else "",
            "clinical_evidence": self.safety_efficacy.clinical_evidence_summary,
            "literature_review": self.safety_efficacy.literature_review_summary,
            "pmcf_information": self.safety_efficacy.pmcf_requirements,
            "benefit_risk_analysis": self.safety_efficacy.risk_benefit_conclusion,
            # New comprehensive fields
            "manufacturer_info": asdict(self.manufacturer),
            "regulatory_info": asdict(self.regulatory),
            "device_identifiers": asdict(self.identifiers),
            "device_detail": asdict(self.device_description),
            "indications_detail": asdict(self.indications),
            "population_info": asdict(self.population),
            "ifu_info": asdict(self.ifu),
            "safety_efficacy_detail": asdict(self.safety_efficacy),
            "tables": self.tables,
            "total_pages": self.total_pages,
            "source_file": self.source_file,
        }
        return flat


def extract_cer_data(filepath: Path) -> CERData:
    """
    Extract comprehensive structured data from a CER document.

    Supports both PDF and DOCX formats. Uses text extraction first,
    then Claude for intelligent structured data extraction.

    Args:
        filepath: Path to the CER file (PDF or DOCX)

    Returns:
        CERData with all extracted information
    """
    filepath = Path(filepath)
    ext = filepath.suffix.lower()

    if ext == ".pdf":
        full_text, tables, page_count = _extract_pdf(filepath)
        extraction_method = "pdf"
    elif ext in (".docx", ".doc"):
        full_text, tables, page_count = _extract_docx(filepath)
        extraction_method = "docx"
    else:
        raise ValueError(f"Unsupported CER format: {ext}. Expected .pdf or .docx")

    # Use Claude to extract structured data from the full text
    cer_data = _ai_extract_cer(full_text, tables)
    cer_data.source_file = filepath.name
    cer_data.total_pages = page_count
    cer_data.extraction_method = extraction_method

    return cer_data


def _extract_pdf(filepath: Path):
    """Extract text and tables from PDF."""
    import pdfplumber

    full_text = ""
    page_texts = []
    tables = []

    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            page_texts.append(text)
            full_text += text + "\n"

            page_tables = page.extract_tables()
            if page_tables:
                for table in page_tables:
                    tables.append({
                        "page": i + 1,
                        "data": table,
                    })

    return full_text, tables, len(page_texts)


def _extract_docx(filepath: Path):
    """Extract text and tables from DOCX."""
    from docx import Document

    doc = Document(str(filepath))
    paragraphs = []
    tables = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            paragraphs.append({"text": text, "style": style})

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append({
                "table_index": i + 1,
                "data": rows,
            })

    full_text = "\n".join(p["text"] for p in paragraphs)

    # Approximate page count from paragraph count
    page_count = max(1, len(paragraphs) // 30)

    return full_text, tables, page_count


def _ai_extract_cer(full_text: str, tables: List[Dict]) -> CERData:
    """Use LLM to extract structured CER data from text and tables."""
    client = get_llm_client()

    # Truncate text if extremely long (keep first and last portions)
    max_chars = 80000
    if len(full_text) > max_chars:
        half = max_chars // 2
        full_text = (
            full_text[:half]
            + "\n\n[... MIDDLE SECTION OMITTED FOR LENGTH ...]\n\n"
            + full_text[-half:]
        )

    # Build tables context
    tables_ctx = ""
    if tables:
        for t in tables[:30]:  # Limit to 30 tables
            location = t.get("page", t.get("table_index", "?"))
            data = t.get("data", [])
            tables_ctx += f"\n--- Table (location: {location}) ---\n"
            for row in data[:20]:  # Limit rows per table
                if row:
                    tables_ctx += " | ".join(str(c) for c in row) + "\n"

    system_prompt = """You are a Clinical Evaluation Report (CER) data extraction specialist.
Extract ALL structured information from the CER document into the specified JSON format.

Rules:
1. Extract ONLY information explicitly present in the document
2. If a field's data is not found in the document, use empty string "" or empty array []
3. Be thorough — CERs contain critical regulatory data scattered across sections
4. For lists (e.g., contraindications, warnings), extract each item separately
5. For tables, describe what data the table contains and which section it belongs to
6. Preserve exact regulatory identifiers (UDI, GMDN codes, certificate numbers)
7. Output valid JSON only, no markdown or explanation
8. CRITICAL: Extract the MANUFACTURER / LEGAL MANUFACTURER name and address. This is typically on the cover page, header, or in the regulatory section.
9. CRITICAL: Extract the EU Technical Documentation (TD) number if mentioned (e.g., TD103, TD-103, Technical Documentation No. 103)
10. CRITICAL: Extract the MDR classification rule (e.g., "Rule 8", "Rule 6") — this is different from the class (IIb).
11. CRITICAL: Extract the Risk Management File (RMF) number if mentioned
12. CRITICAL: Extract the first CE marking date and first Declaration of Conformity date if mentioned
13. CRITICAL: Extract the Authorized Representative (AR) name, address, and SRN"""

    user_prompt = f"""Extract all structured data from this Clinical Evaluation Report.

## CER DOCUMENT TEXT

{full_text}

## EMBEDDED TABLES

{tables_ctx if tables_ctx else "No tables extracted."}

## REQUIRED OUTPUT FORMAT

Return a single JSON object with these top-level keys:

{{
  "manufacturer": {{
    "company_name": "string (legal manufacturer name, e.g. 'CooperSurgical, Inc.')",
    "address_lines": ["string (each line of the manufacturer address)"],
    "manufacturer_srn": "string (SRN if mentioned, e.g. 'US-MF-000002607')",
    "country": "string (manufacturer country)",
    "authorized_representative_name": "string (EU AR name if mentioned)",
    "authorized_representative_address": ["string (AR address lines)"],
    "authorized_representative_srn": "string (AR SRN if mentioned)"
  }},
  "regulatory": {{
    "notified_body_name": "string",
    "notified_body_number": "string (4-digit)",
    "ce_mark_status": "string (active/pending/expired/unknown)",
    "classification_eu": "string (I/IIa/IIb/III or empty)",
    "classification_us_fda": "string (I/II/III or empty)",
    "classification_rule": "string (e.g. 'Rule 8', 'Rule 6' from MDR Annex VIII)",
    "conformity_route": "string (e.g. Annex IX, Annex XI)",
    "certificates": [{{ "type": "string", "number": "string", "date": "string", "expiry": "string" }}],
    "mdr_mdd_reference": "string",
    "regulatory_history": "string (brief timeline of regulatory events)",
    "fda_clearance": "string (510(k) number if mentioned)",
    "eu_technical_documentation_number": "string (e.g. 'TD103', 'TD-103')",
    "first_ce_marking_date": "string (date of first CE marking, ISO format preferred)",
    "first_declaration_of_conformity_date": "string (date of first DoC)",
    "risk_management_file_number": "string (e.g. 'RMF-103', 'RMF103')"
  }},
  "identifiers": {{
    "udi_di": "string",
    "udi_pi": "string",
    "gmdn_codes": [{{ "code": "string", "term": "string" }}],
    "emdn_codes": [{{ "code": "string", "term": "string" }}],
    "model_numbers": ["string"],
    "catalog_numbers": ["string"],
    "manufacturer_identifiers": ["string"]
  }},
  "device_description": {{
    "device_name": "string",
    "device_variants": ["string"],
    "accessories": ["string"],
    "components": ["string"],
    "materials": ["string (e.g. silicone, stainless steel)"],
    "technology": "string (mechanism of action, working principle)",
    "intended_lifespan": "string",
    "sterilization": "string (method if mentioned)",
    "packaging": "string",
    "novel_features": "string"
  }},
  "indications": {{
    "intended_purpose": "string (full intended purpose statement)",
    "medical_indications": ["string"],
    "contraindications": ["string"],
    "warnings": ["string"],
    "precautions": ["string"]
  }},
  "population": {{
    "target_population": "string (general description)",
    "age_range": "string (e.g. adults 18+, pediatric, all ages)",
    "inclusion_criteria": ["string"],
    "exclusion_criteria": ["string"],
    "anatomical_site": "string",
    "clinical_conditions": ["string"],
    "estimated_patient_exposure": "string"
  }},
  "ifu": {{
    "ifu_summary": "string (key IFU content mentioned in CER)",
    "use_instructions": ["string (key steps or procedures)"],
    "cleaning_reprocessing": "string (if reusable device)",
    "storage_handling": "string",
    "training_requirements": "string"
  }},
  "safety_efficacy": {{
    "clinical_evidence_summary": "string (overall summary of clinical evidence)",
    "clinical_investigations": ["string (list of clinical studies referenced)"],
    "literature_review_summary": "string",
    "equivalence_assessment": "string (equivalent device comparison if applicable)",
    "pmcf_requirements": "string",
    "pmcf_planned_activities": ["string"],
    "risk_benefit_conclusion": "string",
    "residual_risks": ["string"],
    "state_of_the_art": "string",
    "overall_conclusions": "string"
  }},
  "tables_summary": [
    {{
      "title": "string (describe what this table contains)",
      "section": "string (which CER section it belongs to)",
      "key_data": "string (summary of the table's key data points)"
    }}
  ]
}}

JSON only:"""

    try:
        response = client.messages.create(
            model=MODEL,
            max_tokens=8192,
            temperature=0.0,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )

        content = response.content[0].text.strip()

        # Strip markdown code fences
        if content.startswith("```"):
            lines = content.split("\n")
            content = "\n".join(
                lines[1:-1] if lines[-1].strip() == "```" else lines[1:]
            )

        # Parse JSON with brace matching for robustness
        extracted = _parse_json_robust(content)

        return _build_cer_data(extracted, tables)

    except Exception as e:
        logger.error(f"AI CER extraction failed: {e}")
        # Return empty CERData — system continues with whatever text extraction got
        return CERData()


def _parse_json_robust(content: str) -> Dict[str, Any]:
    """Parse JSON with fallback brace-matching."""
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass

    # Try brace matching
    brace_depth = 0
    start = content.index("{")
    for i in range(start, len(content)):
        if content[i] == "{":
            brace_depth += 1
        elif content[i] == "}":
            brace_depth -= 1
            if brace_depth == 0:
                return json.loads(content[start : i + 1])

    raise json.JSONDecodeError("No valid JSON object found", content, 0)


def _build_cer_data(extracted: Dict[str, Any], raw_tables: List[Dict]) -> CERData:
    """Build CERData from AI-extracted dict."""

    def _safe_get(d: Dict, key: str, default=None):
        return d.get(key, default) if isinstance(d, dict) else default

    mfr = _safe_get(extracted, "manufacturer", {})
    reg = _safe_get(extracted, "regulatory", {})
    ids = _safe_get(extracted, "identifiers", {})
    desc = _safe_get(extracted, "device_description", {})
    inds = _safe_get(extracted, "indications", {})
    pop = _safe_get(extracted, "population", {})
    ifu_d = _safe_get(extracted, "ifu", {})
    safe = _safe_get(extracted, "safety_efficacy", {})
    tables_summary = _safe_get(extracted, "tables_summary", [])

    cer = CERData()

    # Manufacturer
    cer.manufacturer = ManufacturerInfo(
        company_name=_safe_get(mfr, "company_name", ""),
        address_lines=_safe_get(mfr, "address_lines", []),
        manufacturer_srn=_safe_get(mfr, "manufacturer_srn", ""),
        country=_safe_get(mfr, "country", ""),
        authorized_representative_name=_safe_get(mfr, "authorized_representative_name", ""),
        authorized_representative_address=_safe_get(mfr, "authorized_representative_address", []),
        authorized_representative_srn=_safe_get(mfr, "authorized_representative_srn", ""),
    )

    # Regulatory
    cer.regulatory = RegulatoryInfo(
        notified_body_name=_safe_get(reg, "notified_body_name", ""),
        notified_body_number=_safe_get(reg, "notified_body_number", ""),
        ce_mark_status=_safe_get(reg, "ce_mark_status", ""),
        classification_eu=_safe_get(reg, "classification_eu", ""),
        classification_us_fda=_safe_get(reg, "classification_us_fda", ""),
        classification_rule=_safe_get(reg, "classification_rule", ""),
        conformity_route=_safe_get(reg, "conformity_route", ""),
        certificates=_safe_get(reg, "certificates", []),
        mdr_mdd_reference=_safe_get(reg, "mdr_mdd_reference", ""),
        regulatory_history=_safe_get(reg, "regulatory_history", ""),
        fda_clearance=_safe_get(reg, "fda_clearance", ""),
        eu_technical_documentation_number=_safe_get(reg, "eu_technical_documentation_number", ""),
        first_ce_marking_date=_safe_get(reg, "first_ce_marking_date", ""),
        first_declaration_of_conformity_date=_safe_get(reg, "first_declaration_of_conformity_date", ""),
        risk_management_file_number=_safe_get(reg, "risk_management_file_number", ""),
    )

    # Identifiers
    cer.identifiers = DeviceIdentifiers(
        udi_di=_safe_get(ids, "udi_di", ""),
        udi_pi=_safe_get(ids, "udi_pi", ""),
        gmdn_codes=_safe_get(ids, "gmdn_codes", []),
        emdn_codes=_safe_get(ids, "emdn_codes", []),
        model_numbers=_safe_get(ids, "model_numbers", []),
        catalog_numbers=_safe_get(ids, "catalog_numbers", []),
        manufacturer_identifiers=_safe_get(ids, "manufacturer_identifiers", []),
    )

    # Device Description
    cer.device_description = DeviceDescription(
        device_name=_safe_get(desc, "device_name", ""),
        device_variants=_safe_get(desc, "device_variants", []),
        accessories=_safe_get(desc, "accessories", []),
        components=_safe_get(desc, "components", []),
        materials=_safe_get(desc, "materials", []),
        technology=_safe_get(desc, "technology", ""),
        intended_lifespan=_safe_get(desc, "intended_lifespan", ""),
        sterilization=_safe_get(desc, "sterilization", ""),
        packaging=_safe_get(desc, "packaging", ""),
        novel_features=_safe_get(desc, "novel_features", ""),
    )

    # Indications
    cer.indications = IndicationsInfo(
        intended_purpose=_safe_get(inds, "intended_purpose", ""),
        medical_indications=_safe_get(inds, "medical_indications", []),
        contraindications=_safe_get(inds, "contraindications", []),
        warnings=_safe_get(inds, "warnings", []),
        precautions=_safe_get(inds, "precautions", []),
    )

    # Population
    cer.population = PopulationInfo(
        target_population=_safe_get(pop, "target_population", ""),
        age_range=_safe_get(pop, "age_range", ""),
        inclusion_criteria=_safe_get(pop, "inclusion_criteria", []),
        exclusion_criteria=_safe_get(pop, "exclusion_criteria", []),
        anatomical_site=_safe_get(pop, "anatomical_site", ""),
        clinical_conditions=_safe_get(pop, "clinical_conditions", []),
        estimated_patient_exposure=_safe_get(pop, "estimated_patient_exposure", ""),
    )

    # IFU
    cer.ifu = IFUInfo(
        ifu_summary=_safe_get(ifu_d, "ifu_summary", ""),
        use_instructions=_safe_get(ifu_d, "use_instructions", []),
        cleaning_reprocessing=_safe_get(ifu_d, "cleaning_reprocessing", ""),
        storage_handling=_safe_get(ifu_d, "storage_handling", ""),
        training_requirements=_safe_get(ifu_d, "training_requirements", ""),
    )

    # Safety/Efficacy
    cer.safety_efficacy = SafetyEfficacy(
        clinical_evidence_summary=_safe_get(safe, "clinical_evidence_summary", ""),
        clinical_investigations=_safe_get(safe, "clinical_investigations", []),
        literature_review_summary=_safe_get(safe, "literature_review_summary", ""),
        equivalence_assessment=_safe_get(safe, "equivalence_assessment", ""),
        pmcf_requirements=_safe_get(safe, "pmcf_requirements", ""),
        pmcf_planned_activities=_safe_get(safe, "pmcf_planned_activities", []),
        risk_benefit_conclusion=_safe_get(safe, "risk_benefit_conclusion", ""),
        residual_risks=_safe_get(safe, "residual_risks", []),
        state_of_the_art=_safe_get(safe, "state_of_the_art", ""),
        overall_conclusions=_safe_get(safe, "overall_conclusions", ""),
    )

    # Tables — combine raw extracted tables with AI summaries
    cer.tables = []
    for ts in tables_summary:
        cer.tables.append({
            "title": _safe_get(ts, "title", ""),
            "section": _safe_get(ts, "section", ""),
            "key_data": _safe_get(ts, "key_data", ""),
        })

    return cer
